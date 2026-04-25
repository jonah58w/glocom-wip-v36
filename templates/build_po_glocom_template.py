# -*- coding: utf-8 -*-
"""
產生 GLOCOM 抬頭工廠 PO Word template (v3 - 多品項支援)。

跟 v2 的差別:
- 品項列改用 {%tr for item in items %}...{%tr endfor %} 支援多品項
- 其他版面與 v2 一致(對照鉅盛範本 + 真實附件 ET1150029-01)

執行: python templates/build_po_glocom_template.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).parent
OUTPUT = HERE / "PO_GLOCOM.docx"

FONT_ZH = "標楷體"


def _set_run_font(run, name=FONT_ZH, size=10, bold=False, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _write_cell(cell, text, *, bold=False, size=10, align=None, valign=None,
                line_spacing=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    if line_spacing:
        p.paragraph_format.line_spacing = line_spacing
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
            if align is not None:
                p.alignment = align
        run = p.add_run(line)
        _set_run_font(run, size=size, bold=bold)
    if valign:
        cell.vertical_alignment = valign


def _set_cell_border(cell, top=True, bottom=True, left=True, right=True,
                     sz=4, color="000000"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge, enabled in [
        ('top', top), ('bottom', bottom), ('left', left), ('right', right)
    ]:
        elem = tcBorders.find(qn(f'w:{edge}'))
        if elem is None:
            elem = OxmlElement(f'w:{edge}')
            tcBorders.append(elem)
        if enabled:
            elem.set(qn('w:val'), 'single')
            elem.set(qn('w:sz'), str(sz))
            elem.set(qn('w:space'), '0')
            elem.set(qn('w:color'), color)
        else:
            elem.set(qn('w:val'), 'nil')


def _merge_cells(table, row_start, col_start, row_end, col_end):
    a = table.cell(row_start, col_start)
    b = table.cell(row_end, col_end)
    return a.merge(b)


def build_template():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # ─── 抬頭區 ────────────────────────────────────────
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    for cell in header_table.rows[0].cells:
        _set_cell_border(cell, top=False, bottom=False, left=False, right=False)
    header_table.columns[0].width = Cm(5)
    header_table.columns[1].width = Cm(13)

    left_cell = header_table.rows[0].cells[0]
    _write_cell(left_cell, "GLOCOM", bold=True, size=20,
                align=WD_ALIGN_PARAGRAPH.LEFT)
    p2 = left_cell.add_paragraph()
    r = p2.add_run("[Word 內換成 Logo 圖]")
    _set_run_font(r, size=8)

    right_cell = header_table.rows[0].cells[1]
    right_cell.text = ""
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for txt, sz, bold in [
        ("西 拓 電 子 有 限 公 司\n", 14, True),
        ("GLOCOM ELECTRONICS LTD.\n", 13, True),
        ("3FL.-2, NO.268, LIANCHENG RD., ZHONGHE\n", 9, False),
        ("NEW TAIPEI CITY 23553, TAIWAN R.O.C.\n", 9, False),
        ("新北市中和區連城路268號3樓之2; 統一編號:70489361\n", 9, False),
        ("PHONE: 886 2 82273189; FAX: 886 2 82273187", 9, False),
    ]:
        run = p.add_run(txt)
        _set_run_font(run, size=sz, bold=bold)

    # ─── REVISED 標記 ───────────────────────────────────
    rp = doc.add_paragraph()
    rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = rp.add_run("{% if order.is_revised %}[ REVISED ]{% endif %}")
    _set_run_font(r, size=14, bold=True, color=RGBColor(0xC0, 0x00, 0x00))

    # ─── 主表格 ─────────────────────────────────────────
    # Row 0: 標題;Row 1-4: 廠商/採購;Row 5: 品項表頭;Row 6: 品項列(for-loop)
    # Row 7: 合計;Row 8: 付款方式;Row 9: 注意事項標題;Row 10-19: 條款;Row 20-21: 簽章
    main_table = doc.add_table(rows=22, cols=6)
    main_table.style = "Table Grid"
    main_table.autofit = False

    widths = [Cm(2.4), Cm(4.6), Cm(2.5), Cm(2.4), Cm(2.4), Cm(2.9)]
    for i, w in enumerate(widths):
        main_table.columns[i].width = w
        for cell in main_table.columns[i].cells:
            cell.width = w

    # Row 0
    title_cell = _merge_cells(main_table, 0, 0, 0, 5)
    _write_cell(title_cell, "訂  購  單", bold=True, size=16,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                valign=WD_ALIGN_VERTICAL.CENTER)

    # Row 1-4
    rows_left_right = [
        ("廠商名稱:{{ factory.name }}", "頁    次:1/1"),
        ("廠商地址:{{ factory.address }}", "採購單號:{{ order.no }}"),
        ("聯  絡  人:{{ factory.contact }}", "採購日期:{{ order.date_display }}"),
        (
            "電    話:{{ factory.phone }}     傳    真:{{ factory.fax }}",
            "負責採購:{{ order.purchase_responsible }}\n幣    別:{{ order.currency }}",
        ),
    ]
    for i, (left_text, right_text) in enumerate(rows_left_right):
        row_idx = 1 + i
        left = _merge_cells(main_table, row_idx, 0, row_idx, 3)
        _write_cell(left, left_text, size=10, valign=WD_ALIGN_VERTICAL.CENTER)
        right = _merge_cells(main_table, row_idx, 4, row_idx, 5)
        _write_cell(right, right_text, size=10, valign=WD_ALIGN_VERTICAL.CENTER)

    # Row 5: 品項表頭
    headers = ["產品編號", "產品規格", "交期", "數量", "單價", "小計"]
    for i, h in enumerate(headers):
        _write_cell(
            main_table.cell(5, i), h, bold=True, size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            valign=WD_ALIGN_VERTICAL.CENTER,
        )

    # Row 6: 品項內容(單品項版,實際多品項由 Python 端複製 row)
    # 這樣比跟 docxtpl {%tr for %} 對抗簡單,而且行為可預測
    _write_cell(main_table.cell(6, 0), "{{ item.part_number }}",
                size=10, valign=WD_ALIGN_VERTICAL.TOP)
    _write_cell(main_table.cell(6, 1), "{{ item.spec_text }}",
                size=10, valign=WD_ALIGN_VERTICAL.TOP)
    _write_cell(main_table.cell(6, 2), "{{ item.delivery_display }}",
                size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                valign=WD_ALIGN_VERTICAL.TOP)
    _write_cell(main_table.cell(6, 3), "{{ item.quantity_display }}",
                size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
                valign=WD_ALIGN_VERTICAL.TOP)
    _write_cell(main_table.cell(6, 4), "{{ item.unit_price }}",
                size=10, align=WD_ALIGN_PARAGRAPH.RIGHT,
                valign=WD_ALIGN_VERTICAL.TOP)
    _write_cell(main_table.cell(6, 5), "{{ item.amount }}",
                size=10, align=WD_ALIGN_PARAGRAPH.RIGHT,
                valign=WD_ALIGN_VERTICAL.TOP)

    main_table.rows[6].height = Cm(2.5)

    # Row 7: 合計
    total_label = _merge_cells(main_table, 7, 0, 7, 4)
    _write_cell(total_label, "合      計", bold=True, size=11,
                align=WD_ALIGN_PARAGRAPH.RIGHT,
                valign=WD_ALIGN_VERTICAL.CENTER)
    _write_cell(
        main_table.cell(7, 5),
        "{{ order.currency }} {{ order.total_amount }}",
        bold=True, size=11,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        valign=WD_ALIGN_VERTICAL.CENTER,
    )

    # Row 8: 付款方式 / 進貨方式 / 進貨地址
    payment_cell = _merge_cells(main_table, 8, 0, 8, 5)
    _write_cell(
        payment_cell,
        "付款方式:{{ order.payment_terms }}\n"
        "進貨方式:{{ order.shipment_method }}\n"
        "進貨地址:{{ order.ship_to }}",
        size=10, valign=WD_ALIGN_VERTICAL.TOP,
    )

    # Row 9: 注意事項標題
    notes_title = _merge_cells(main_table, 9, 0, 9, 5)
    _write_cell(notes_title, "注意事項:", bold=True, size=10,
                valign=WD_ALIGN_VERTICAL.CENTER)

    # Row 10-19: 注意事項 10 條(您現行版本)
    NOTES = [
        "出貨時須附 出貨報告、C of C、切片報告(含切片)& 試錫板(真空包裝)。",
        "真空包裝每包的毛重不可超過19Kgs;而板厚1.6mm以下,15PNL一包;1.6mm-2.0mm(不含,以下),10PNL一包;2.0mm(含以上),5PNL一包。",
        "須符合客戶SPECs與IPC6012規定,不可補線,但經事先允許者除外。",
        "出貨報告 & C of C:用GLOCOM抬頭文件,且須用西拓UL Logo,除非不需UL。(請勿在箱內放置任何不符規定之貴司名稱的文件,若致各種異常或違背一般商業保密協定,則責任自負!)",
        "須距出貨交期一個月內才可發料生產,否則,因修改、變更REV或Cancel訂單所造成之損失,須自行負責。",
        "當訂單與所提供之Gerbers & drawing有衝突不同時(如料號或版本等),以訂單為主,但仍須提出詢問確認後,方可生產,否則須自行負責所造成之錯誤。",
        "若出庫存板,務必注意Date Code,化金/HASL不能超過五個月,其餘表面處理不能超過二個月,否則客戶會剔退板子。",
        "務必使用外銷箱(厚度:7-8mm以上);箱內四周與上下面,須內襯厚紙板,並以瓦楞紙板或氣泡墊(不可使用保麗龍!)塞滿空隙,再以打包帶打包,除正反兩面外,中間須再固定一條,以免搬運過程中有箱子破裂,或板子擠壓等問題發生!",
        "出貨板子,不可有MADE IN USA或其他國家;否則一經海關驗出,須自負罰責。",
        "若從大陸出貨,所有與物流Booking的文件,寄件人須填寫\"費司名稱/(Eusway/Glocom)\",若未加註(Eusway/Glocom),致客戶端IQC進料檢驗拒收,須自行負責。",
    ]
    for i, note in enumerate(NOTES):
        row_idx = 10 + i
        merged = _merge_cells(main_table, row_idx, 0, row_idx, 5)
        _write_cell(merged, f"{i+1}. {note}", size=9,
                    valign=WD_ALIGN_VERTICAL.TOP,
                    line_spacing=1.15)

    # Row 20: 簽章標題
    sig_left_label = _merge_cells(main_table, 20, 0, 20, 2)
    _write_cell(sig_left_label, "廠商確認:\n{{ factory.name }}",
                size=10, valign=WD_ALIGN_VERTICAL.TOP)
    sig_right_label = _merge_cells(main_table, 20, 3, 20, 5)
    _write_cell(sig_right_label, "公司確認:\n西 拓 電 子 有 限 公 司",
                size=10, valign=WD_ALIGN_VERTICAL.TOP)

    # Row 21: 簽名空白
    sig_left_blank = _merge_cells(main_table, 21, 0, 21, 2)
    _set_cell_border(sig_left_blank, top=False, left=False, right=False)
    _write_cell(sig_left_blank, "\n\n____________________________",
                size=10, valign=WD_ALIGN_VERTICAL.BOTTOM)

    sig_right_blank = _merge_cells(main_table, 21, 3, 21, 5)
    _set_cell_border(sig_right_blank, top=False, left=False, right=False)
    _write_cell(sig_right_blank, "\n\n____________________________",
                size=10, valign=WD_ALIGN_VERTICAL.BOTTOM)

    doc.save(OUTPUT)
    print(f"✓ Template generated: {OUTPUT}")
    return OUTPUT


if __name__ == "__main__":
    build_template()
