# -*- coding: utf-8 -*-
"""
fix_template.py — 自動修補 PO_GLOCOM.docx 的多品項迴圈 bug。

問題:
  原本 templates/PO_GLOCOM.docx 的品項 row 只有 {{ item.xxx }} placeholder,
  沒有 {%tr for item in items %} 跟 {%tr endfor %} 標籤,
  導致 docxtpl 渲染時只展開第一筆 item,後續品項被吞掉。

修法:
  在「品項 row」(Table 2, Row 1) 前面插入一個只含 {%tr for item in items %} 的 row,
  後面插入一個只含 {%tr endfor %} 的 row。
  這兩個 tag-only row 在 docxtpl 預處理時會被消化(整個 <w:tr> 替換成 jinja2 tag),
  剩下品項 row 被 jinja2 for 迴圈展開,每筆 item 一行。

關鍵:for/endfor 必須在「不同的 <w:tr>」裡。
  如果寫在同一個 <w:tr> 兩個 cell,docxtpl 的 regex 會把整個 row 替換成第一個 tag,
  endfor 會跟 row 一起被消化掉,造成 jinja2 報「unknown tag 'endfor'」。

使用:
  python fix_template.py
  → 讀 templates/PO_GLOCOM.docx,輸出修補後的同檔。
"""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


HERE = Path(__file__).parent
TEMPLATE = HERE / "PO_GLOCOM.docx"


def make_tag_only_row(template_row_element, tag_text: str):
    """
    複製 template_row_element 的結構,清空所有 cell text,
    只在第一個 cell 寫入 tag_text。
    """
    new_row = deepcopy(template_row_element)
    cells = new_row.findall(qn("w:tc"))
    for c_idx, cell in enumerate(cells):
        # 清空 cell 的所有 text
        for p in cell.findall(qn("w:p")):
            for r in p.findall(qn("w:r")):
                for t in r.findall(qn("w:t")):
                    t.text = ""
        if c_idx == 0:
            ps = cell.findall(qn("w:p"))
            if ps:
                rs = ps[0].findall(qn("w:r"))
                if rs:
                    ts = rs[0].findall(qn("w:t"))
                    if ts:
                        ts[0].text = tag_text
                    else:
                        t = etree.SubElement(rs[0], qn("w:t"))
                        t.text = tag_text
                else:
                    r = etree.SubElement(ps[0], qn("w:r"))
                    t = etree.SubElement(r, qn("w:t"))
                    t.text = tag_text
    return new_row


def already_patched(items_table) -> bool:
    """檢查模板是否已經有 tr-loop tag(避免重複修補造成 row 越來越多)"""
    for row in items_table.rows:
        for cell in row.cells:
            text = cell.text
            if "{%tr for item" in text or "{% tr for item" in text:
                return True
    return False


def fix_template(template_path: Path):
    if not template_path.exists():
        raise FileNotFoundError(f"找不到 {template_path}")

    doc = Document(str(template_path))

    # Table 2 是品項表(Table 0=抬頭, 1=廠商資訊, 2=品項+合計, 3=簽章)
    if len(doc.tables) < 3:
        raise RuntimeError(
            f"模板結構不符預期(只有 {len(doc.tables)} 個 tables,預期至少 3 個)"
        )

    items_table = doc.tables[2]
    if already_patched(items_table):
        print("✓ 模板已有 tr-loop tag,不需要再修補")
        return

    if len(items_table.rows) < 2:
        raise RuntimeError(
            f"Table 2 結構不符(只有 {len(items_table.rows)} rows,預期至少有表頭+品項+合計)"
        )

    # Row 1 是品項 placeholder row
    item_row_element = items_table.rows[1]._tr

    # 在 row 1 前插入 for-tag-only row,後插入 endfor-tag-only row
    for_row = make_tag_only_row(item_row_element, "{%tr for item in items %}")
    endfor_row = make_tag_only_row(item_row_element, "{%tr endfor %}")

    item_row_element.addprevious(for_row)
    item_row_element.addnext(endfor_row)

    doc.save(str(template_path))
    print(f"✓ 修補完成:{template_path}")
    print("  - 在品項 row 前面插入了 {%tr for item in items %} 的 tag-row")
    print("  - 在品項 row 後面插入了 {%tr endfor %} 的 tag-row")


if __name__ == "__main__":
    fix_template(TEMPLATE)
