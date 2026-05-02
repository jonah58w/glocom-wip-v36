# -*- coding: utf-8 -*-
"""
建立工廠 PO 頁面 v3.9.19。

主要更新(v3.9.19):
- ★ 全面強化「歷史工廠面板」的型別防禦,任何意外都不讓 page 崩:
       1. spec_history_data 不是 dict → 顯示 warning + fallback 空 dict
       2. parsed_items 元素沒 .part_number → 用 getattr 安全跳過
       3. record 的 .get() 炸 → try/except 跳過
       4. 整個面板渲染包 try/except → 失敗只 show error 不 break page
       5. _aggregate 也加防禦,suggestions 異常時回傳空 Counter

v3.9.18 變更:
- 對 record 做 isinstance(dict) 防禦,跳過非 dict 元素

v3.9.14 變更:
- 產 PDF 後記錄並顯示「這份 PDF 是用哪家工廠/字首/編號產的」
- Step 6 加「PDF 即將使用的工廠資料」預覽
- Step 5 default 來源資訊改放最終規格上方

v3.9.13 變更:
- 切工廠不再把單價歸零(sentinel 簡化)
- Step 4 / Step 6 加 debug expander

v3.9.12 變更:
- 修「規格內容重複舊料號」bug
- 工廠單價自動抓前批(同字首+同工廠優先)

v3.9.11 變更:
- 歷史 default 也按字首(GC/ET/EW)抓
- 歷史工廠 expander 每筆加字首 badge
- NRE 真的預設 0(每次換 PDF 強制重置)

v3.9.10 變更:
- 修正:final_key 為空時自動重新套 default
- 修正:Step 3 切工廠時,Step 5 自動用「同工廠最近一筆」當 default
- 新功能:「規格變更」text_input 加 on_change callback,按 Enter 自動套用
- 改進:歷史工廠 expander 永遠預設展開

v3.9.9 變更:
- PDF 排版修正:在每個「數字. 」開頭的行前面加零寬空格 \\u200B,
       避免 Word/LibreOffice 把它認成 numbered list 自動套 tab stop
- 新功能:Step 5 舊料號區塊加「🏭 歷史工廠」expander,
       列出此料號歷史做過的所有工廠 + 各自規格,可一鍵切換
- NRE 金額預設改成 "0"

v3.9.8 變更:
- 修正:歷史 spec_text 載入時自動清洗 tab(去掉 "1.\t..." 對齊用 tab),
       印到 PDF 不再有編號項目後的大空白
- 新功能:Step 5 舊料號區塊加「📝 規格變更」欄位
       - 用在「舊料號但這次規格有變更」(譬如表面處理改 ENIG 2u")
       - 按「⬇ 套用變更」會把變更插到歷史規格頂端「舊料號」後
       - 寫回 Teable 後也會更新 spec_history.json,下次同料號自動帶新版

v3.9.7 變更:
- 新功能:Step 6 加「📑 產 PI」按鈕,產 Proforma Invoice
- 用 templates/PI_GLOCOM.docx 模板 + core/pi_generator.py
- 客戶資訊從 data/customers.json 讀(VORNE/WESCO 等)
- Invoice No. = 主號(去 -01),譬如 G1150034-01 → G1150034
- 品項 = PCB 主品項 + Setup & Tooling Charge + Bank fee
- Bank fee 預設 $45,Sandy 可改

v3.9.6 變更:
- Teable 寫入訊息覆蓋 + 偵錯資訊

v3.9.5 變更:
- ★ NRE 註記新格式 — 含原單價,讓工廠看到完整計算:
  範例:
    130.000
    (內含
    單價
    NTD102/pcs
    + NRE:
    NTD14,000;
    NTD28/pcs)
  工廠一眼就懂:130 = 102(原 PCB)+ 28(NRE 攤提)
- 主品項合併前的「原工廠單價」存進 nre_settings.original_unit_price

v3.9.4 變更:
- NRE 註記移到「單價欄底下」(原本在規格欄末尾)
- docx 模板欄寬調整:單價 1.7cm → 2.5cm,規格欄保持 5.5cm

v3.9.3 變更:
- NRE 註記改用「(內含 NRE: NTD14,000; NTD28/pcs)」格式
- NT$ → NTD 標準化
- 修 docx 模板 `&` 消失 bug

v3.9.2 變更:
- 「新料號」介面重新設計 — 7 選單預設展開,不再藏在 tab
- NRE 偵測規則收緊(料號 99-開頭 + 描述開頭 NRE/Setup Fee 等,數量=1)

v3.9 變更:
- 新料號規格 3 段式組合(Working Gerber 確認 / 7 選單 / 結尾標準語)
- Surface Finish: ENIG 2u" 取代重複的 ENIG/Immersion Gold

v3.8.1 變更:
- NRE 工程費獨立欄位 + 工廠貨幣自動 (NTD/USD)
- 字首決定處理模式(GC=合併, ET/EW=獨立)
- 填 0/NA/留空 → 跳過 NRE

v3.8 變更:
- NRE 工程費自動處理初版

v3.7 變更:
- ★ AI 規格智能分析(在 Step 5 舊料號區塊):
  - 模糊料號比對(ATP3 Rev G == ATP3-Rev-G == atp3 rev g)
  - 行頻率分析:識別「核心要求 / 最新加 / 舊版有最新沒」
  - 自動產生「智能合併版本」(以最新為主 + 標出舊版疑點)
  - 顯示統計徽章 + 紅色警告 + 詳細分析展開
  - Sandy 看高亮疑點檢查,不會漏掉重要要求

v3.6 變更:
- 寫進 Teable 成功後,自動把規格寫進 GitHub 上的 spec_history.json
  - 透過 GitHub REST API,不依賴本機 git
  - PAT 存在 Streamlit Secrets ([github] token = "...")
  - 兩條更新管道並行(ERP 排程 + Streamlit 即時)

v3.5 變更:
- 三層編號防撞機制
- Refresh 按鈕

v3.4 變更:
- 規格優先從 data/spec_history.json(歷史 RTF/DOCX)抓
- 找不到再從 Teable 主表「客戶要求注意事項」抓

v3.3 變更:
- 規格頂部加「舊料號 / 新料號」radio
- 選舊料號自動帶歷史注意事項

v3.2 變更:
- 工廠下拉用 sort_priority 排序

v3.1 變更:
- 拿掉「規格沒填→fallback 客戶描述」邏輯,避免規格欄重複印料號
- 規格沒填時顯示紅色錯誤,阻擋產 PDF / 寫回 Teable
"""

from __future__ import annotations
import json
import re
import traceback
from collections import Counter
from datetime import date, datetime
from pathlib import Path

import pandas as pd
import requests
import streamlit as st

from core.customer_po_parser import (
    parse_customer_po_from_pdf,
    POItem,
    ParsedPO,
)
from core.teable_query import (
    parse_glocom_po_no,
    safe_text,
    COL_GLOCOM_PO,
)


# ─── 設定 ────────────────────────────────────────
HERE = Path(__file__).parent
FACTORIES_JSON = HERE / "data" / "factories.json"
SPEC_HISTORY_JSON = HERE / "data" / "spec_history.json"


# 新料號完整規格的選項
SPEC_OPTIONS_FULL = {
    "Material": ["2L", "4L", "6L", "8L", "10L", "Aluminum", "FR4", "Rogers"],
    "Tg": ["Tg130", "Tg150", "Tg170", "Tg180"],
    "Board thickness": ["0.4mm", "0.6mm", "0.8mm", "1.0mm", "1.2mm", "1.6mm", "2.0mm", "2.4mm", "3.2mm"],
    "Copper": ["1oz/1oz", "1/2oz/1/2oz", "2oz/2oz", "3oz/3oz", "Ext: 1oz all layers"],
    # ★ v3.9: Surface Finish 整理 — ENIG 跟 Immersion Gold 同一個意思,只留 ENIG 2u"
    "Surface Finish": ['ENIG 2u"', "Lead-free HASL", "HASL", "OSP", "Hard Gold", "Immersion Silver"],
    "S/M": ["Green", "Matte Green", "Red", "Blue", "Black", "Matte Black", "White"],
    "S/L": ["White", "Black", "Yellow"],
}

# ★ v3.9: 新料號規格 - 第 1 段(工程確認語)選項
# 字串保持乾淨 &(顯示給 Sandy),寫進 docx 模板前才會 escape 成 &amp;
ENGINEERING_CONFIRM_OPTIONS = {
    "Working Gerber 承認": (
        "Working Gerber承認後,才可生產! 請於下午14:00前傳working gerber"
        "(& stencil gerber)給我司,如有工程問題請一併詢問,謝謝!"
    ),
    "直接生產": "直接生產!",
    "不寫": "",
}

# ★ v3.9: 新料號規格 - 第 3 段(結尾標準語,強制加,兩行)
FACTORY_PO_FOOTER = (
    "須添加西拓UL logo & date code (YYWW);\n"
    "樣板: 除試錫板外,須另外提供樣板供備份."
)


def escape_for_docx(text: str) -> str:
    """
    把字串裡的 XML 特殊字元 escape,讓 docxtpl |safe filter 能正確渲染。
    主要處理:& → &amp;, < → &lt;, > → &gt;
    \n 不用處理(docxtpl 會自動轉成 <w:br/>)
    """
    if not text:
        return ""
    # 順序很重要:先 & 再其他,否則 &lt; 會變 &amp;lt;
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


# ─── ★ v3.9.8: spec_text 清洗 helper ────────────────────────
def _normalize_spec_text(text: str) -> str:
    """
    清洗 spec_text 裡的 tab 與多餘空白:
    - "1.\t板子不可有..." → "1. 板子不可有..."  (去掉編號項目後的對齊 tab)
    - 行內多個連續空白合成一個
    - 保留換行 \\n(行與行的分隔不動)

    歷史 RTF/DOCX 解析時很多項目是用 tab 對齊("1.\\t..." / "•\\t..."),
    不清掉的話印到 PDF 會留下大塊空白,看起來像被 indent 出去。
    """
    if not text:
        return text
    # tab → 單一空格
    text = text.replace("\t", " ")
    # 行內多個連續空白合成單一空白(換行不動)
    lines = text.split("\n")
    normalized = [re.sub(r" +", " ", line).rstrip() for line in lines]
    return "\n".join(normalized)


def _add_zwsp_to_numbered_lines(text: str) -> str:
    """
    ★ v3.9.9: 在「數字. 」開頭的行前面加零寬空格 \\u200B,
    避免 Word/LibreOffice 把它認成 numbered list 起始,
    導致段落自動套上 list paragraph 樣式 + hanging indent + tab stop
    (這正是規格欄出現「1.   板子...」大空白的元凶 — 即使 source 沒 tab,
     渲染器自己會加)。

    零寬空格在 PDF/Word 上完全看不到,但段落第一個字元變成 ZWSP
    就不會觸發 list 偵測。
    """
    if not text:
        return text
    # 用 lambda 避免 re.sub replacement string 的 \\u200B escape 問題
    return re.sub(
        r"^(\s*)(\d+\. )",
        lambda m: m.group(1) + "\u200B" + m.group(2),
        text,
        flags=re.MULTILINE,
    )


def _merge_spec_with_change(history_spec: str, change_text: str) -> str:
    """
    把 Sandy 這次填的「規格變更」插入到歷史規格的第一行「舊料號」後面。

    case 1: 第一行剛好 == "舊料號"      → "舊料號; {change}"
    case 2: 第一行 "舊料號;..." 或 "舊料號:..." → "舊料號; {change}; {原內容}"
    case 3: 第一行不是「舊料號」開頭   → "舊料號; {change}\\n{原內容}"

    說明:這個函式永遠用「歷史規格原版」當底重組,不會在現況上 append,
        所以 Sandy 重複按「⬇ 套用變更」也不會 double 上去。
    """
    change = (change_text or "").strip().rstrip(";；,，。 ").lstrip(";；,， ")
    history = (history_spec or "").strip()
    if not change:
        return history
    if not history:
        return f"舊料號; {change}"

    lines = history.split("\n")
    first_line = lines[0].rstrip()

    # case 2: 「舊料號; XXX」/「舊料號: XXX」/「舊料號;XXX」/「舊料號;XXX」
    m = re.match(r"^(舊料號)\s*([;；:：])\s*(.*)$", first_line)
    if m:
        rest = m.group(3).strip()
        if rest:
            lines[0] = f"舊料號; {change}; {rest}"
        else:
            lines[0] = f"舊料號; {change}"
        return "\n".join(lines)

    # case 1: 純「舊料號」單獨一行
    if first_line == "舊料號":
        lines[0] = f"舊料號; {change}"
        return "\n".join(lines)

    # case 3: 第一行不是舊料號開頭 → prepend 新的一行
    return f"舊料號; {change}\n{history}"


def _ensure_starts_with_jiu_liao(spec_text: str) -> str:
    """
    ★ v3.9.12: 確保 spec_text 第一行以「舊料號」開頭。
    - 第一行已是「舊料號」開頭(包含「舊料號」/「舊料號;...」/「舊料號;...」/...)→ 原樣回傳
    - 第一行不是 → prepend「舊料號; 」

    這樣可以避免之前 default_text 加 "舊料號\\n\\n注意:\\n" 前綴
    跟歷史 spec_text 本身第一行的「舊料號;...」重複,造成 PO 印兩個「舊料號」。
    """
    text = (spec_text or "").lstrip()
    if not text:
        return "舊料號"
    first_line = text.split("\n", 1)[0]
    if first_line.startswith("舊料號"):
        return text
    return f"舊料號; {text}"


def _po_prefix(po_no: str) -> str:
    """
    ★ v3.9.11: 從 PO 編號取字首類型。
    - "G1150030-01" → "GC"
    - "ET1150009-01" → "ET"
    - "EW1140005-01" → "EW"
    - 空字串或不認得 → ""

    用來:
    1. default 套用時依「同字首」優先排序
    2. 歷史工廠 expander 每筆顯示 GC/ET/EW badge
    """
    po = (po_no or "").strip().upper()
    if po.startswith("ET"):
        return "ET"
    if po.startswith("EW"):
        return "EW"
    if po.startswith("G"):
        return "GC"
    return ""


def _compute_old_pn_factory_suggestions(parsed_items, spec_history_data) -> dict:
    """
    ★ v3.9.16: 對每個品項統計 GC/ET/EW 三個字首下的歷史工廠。
    ★ v3.9.18: 對 record 做 isinstance(dict) 防禦,避免舊格式造成 AttributeError。
    ★ v3.9.19: 連 spec_history_data 本身、parsed_items 元素都做防禦,
       任何意外型別都安全跳過,不讓整個 page 崩。

    回傳:
        {
          P/N: {
            "by_prefix": {"GC": Counter({factory: count}), "ET": ..., "EW": ...},
            "latest": {"GC": (factory, po, date), ...},
            "total": int,
          },
          ...
        }
    """
    suggestions = {}
    # ★ v3.9.19: spec_history_data 本身必須是 dict
    if not isinstance(spec_history_data, dict):
        return suggestions

    if not parsed_items:
        return suggestions

    for item in parsed_items:
        # ★ v3.9.19: 用 getattr 取 part_number,沒這個屬性就跳過
        pn = getattr(item, "part_number", None)
        if not pn or not isinstance(pn, str):
            continue
        pn = pn.strip()
        if not pn:
            continue

        records_raw = spec_history_data.get(pn, [])
        # ★ v3.9.18: records_raw 型別檢查
        if isinstance(records_raw, dict):
            records = [records_raw]
        elif isinstance(records_raw, list):
            records = records_raw
        else:
            records = []

        by_prefix = {"GC": Counter(), "ET": Counter(), "EW": Counter()}
        latest = {"GC": None, "ET": None, "EW": None}

        for r in records:
            # ★ v3.9.18: 跳過非 dict record
            if not isinstance(r, dict):
                continue
            try:
                factory = (r.get("factory") or "").strip() or "(未填工廠)"
                po = (r.get("po") or "").strip()
                date_str = (r.get("date") or "").strip()
            except Exception:
                # 連 .get 都炸的奇怪 case 也跳過
                continue
            prefix = _po_prefix(po)
            if prefix in by_prefix:
                by_prefix[prefix][factory] += 1
                if latest[prefix] is None or date_str > (latest[prefix][2] or ""):
                    latest[prefix] = (factory, po, date_str)

        suggestions[pn] = {
            "by_prefix": by_prefix,
            "latest": latest,
            "total": sum(c.total() for c in by_prefix.values()),
        }
    return suggestions


def _aggregate_factory_prefix_votes(suggestions: dict) -> Counter:
    """
    把每個品項的歷史工廠+字首組合票數加總。
    回傳 Counter[(factory, prefix)] = total_count

    ★ v3.9.19: 加防禦 — suggestions 異常時回傳空 Counter,不讓 panel 崩。
    """
    votes = Counter()
    if not isinstance(suggestions, dict):
        return votes
    for pn, sug in suggestions.items():
        if not isinstance(sug, dict):
            continue
        by_prefix = sug.get("by_prefix", {})
        if not isinstance(by_prefix, dict):
            continue
        for prefix, factories in by_prefix.items():
            if not isinstance(factories, Counter):
                continue
            for factory, count in factories.items():
                if factory == "(未填工廠)":
                    continue
                votes[(factory, prefix)] += count
    return votes


def _on_spec_change_apply(history_default_key: str, change_key: str, final_key: str):
    """
    ★ v3.9.10: 「規格變更」text_input 的 on_change callback。
    Sandy 按 Enter / 失焦時自動觸發,不用再點按鈕。

    從 history_default_key(原版規格)重組,避免重複按時 double 套上去。
    """
    change_text = (st.session_state.get(change_key, "") or "").strip()
    if not change_text:
        return  # 空字串就不動,讓 Sandy 自己控制
    history_default = st.session_state.get(history_default_key, "") or ""
    st.session_state[final_key] = _merge_spec_with_change(history_default, change_text)


# ─── 工廠主檔 ────────────────────────────────────
def _strip_placeholder_brackets(text: str) -> str:
    """剝掉字串中的 [請補...] 方框註記,例如『優技電子 [請補完整名]』→『優技電子』"""
    if not text:
        return ""
    return re.sub(r"\s*\[[^\]]*\]\s*", "", str(text)).strip()


def load_factories() -> dict:
    if not FACTORIES_JSON.exists():
        return {}
    try:
        with open(FACTORIES_JSON, "r", encoding="utf-8") as f:
            data = json.load(f)
        all_factories = data.get("factories", {})
        active = {k: v for k, v in all_factories.items() if v.get("is_active")}
        # 把 factory_name 的方框註記剝掉(避免印在簽名欄)
        for v in active.values():
            v["factory_name"] = _strip_placeholder_brackets(v.get("factory_name", ""))
        return active
    except Exception as e:
        st.error(f"factories.json 讀取失敗:{e}")
        return {}


# ─── 字首邏輯 ────────────────────────────────────
def derive_prefix(issuing_company: str, factory_region: str) -> str:
    issuing = (issuing_company or "").upper()
    region = (factory_region or "").lower()
    if issuing == "GLOCOM":
        return "G"
    if issuing == "EUSWAY":
        if region in ("china", "hk", "hong kong"):
            return "EW"
        return "ET"
    return "G"


def display_prefix(internal_prefix: str) -> str:
    return "GC" if internal_prefix == "G" else internal_prefix


def internal_prefix(display: str) -> str:
    return "G" if display == "GC" else display


# ─── ★ v3.5: 即時撈 Teable API ─────────────
def fetch_all_po_numbers_from_teable(table_url: str, headers: dict) -> set[str]:
    """
    直接打 Teable API 撈所有訂單編號(不靠側邊欄 Refresh 的 cache)。
    
    回傳:set of all 西拓訂單編號 字串
    """
    po_numbers = set()
    page_token = None
    pn = COL_GLOCOM_PO  # "西拓訂單編號"
    
    try:
        while True:
            params = {
                "fieldKeyType": "name",
                "cellFormat": "text",
                "take": 1000,
                "viewId": "viwGqqwdvy5WkaqfQuJ",  # 「API - All Records」
                "projection": [pn],  # 只撈這一欄省流量
            }
            if page_token:
                params["nextPageToken"] = page_token
            
            r = requests.get(table_url, headers=headers, params=params, timeout=30)
            if r.status_code != 200:
                break
            
            data = r.json()
            records = data.get("records", []) or []
            for rec in records:
                fields = rec.get("fields", {}) or {}
                val = str(fields.get(pn, "") or "").strip()
                if val:
                    po_numbers.add(val)
            
            page_token = data.get("nextPageToken")
            if not page_token:
                break
    except Exception as e:
        st.warning(f"⚠️ 即時撈 Teable 編號失敗:{e}(改用側邊欄快取)")
    
    return po_numbers


def calc_next_po_number(po_numbers_set: set[str], prefix: str) -> str:
    """
    計算下一個訂單編號。
    
    重點:流水號要**全局唯一**,不論字首 G/ET/EW/GC。
    例如 Teable 上有 G1150030-01,則下一個 ET 編號也不能用 1150030,
    必須是 ET1150031-01。
    
    Args:
        po_numbers_set: 所有現有訂單編號的 set
        prefix: 內部字首 (G / ET / EW)
    
    Returns: 下一個訂單編號,例如 G1150031-01
    """
    if not po_numbers_set:
        roc_year = datetime.now().year - 1911
        return f"{prefix}{roc_year}0001-01"

    pattern = re.compile(r"^[A-Z]+(\d{3})(\d{4})-(\d+)$")
    max_year_serial = (0, 0)

    for po in po_numbers_set:
        m = pattern.match(po.strip())
        if m:
            year = int(m.group(1))
            serial = int(m.group(2))
            if (year, serial) > max_year_serial:
                max_year_serial = (year, serial)

    if max_year_serial == (0, 0):
        roc_year = datetime.now().year - 1911
        return f"{prefix}{roc_year}0001-01"

    year, serial = max_year_serial
    return f"{prefix}{year}{serial + 1:04d}-01"


def calc_next_po_number_from_df(orders_df: pd.DataFrame, prefix: str) -> str:
    """從 DataFrame (側邊欄快取)算下一編號 - fallback 用"""
    if orders_df.empty or COL_GLOCOM_PO not in orders_df.columns:
        roc_year = datetime.now().year - 1911
        return f"{prefix}{roc_year}0001-01"
    po_set = set()
    for raw_po in orders_df[COL_GLOCOM_PO].dropna().astype(str):
        po_set.add(raw_po.strip())
    return calc_next_po_number(po_set, prefix)


# ─── Teable 寫入 ─────────────────────────────────
def fetch_previous_factory_price(
    orders: pd.DataFrame,
    part_number: str,
    factory_short: str = "",
    po_prefix: str = "",
) -> float | None:
    """
    ★ v3.9.12: 從 Teable 主表查同 P/N 最近一筆的工廠單價。

    計算方法:銷貨金額 ÷ Order Q'TY (PCS)
    排序:同字首 +10, 同工廠 +5, 再按工廠下單日期降冪。

    回傳:單價(float, 保留 4 位小數)或 None(沒歷史)
    """
    if orders is None or orders.empty:
        return None

    pn_col = "P/N"
    if pn_col not in orders.columns:
        return None

    # 找數量欄位(可能是 Order Q'TY\n (PCS) 或 Q'TY (PCS) 等)
    qty_col_candidates = [
        "Order Q'TY\n (PCS)", "Order Q'TY (PCS)", "Order QTY (PCS)",
        "Q'TY (PCS)", "Q'TY", "Qty", "數量", "PCS",
    ]
    qty_col = next((c for c in qty_col_candidates if c in orders.columns), None)
    if not qty_col:
        return None

    amt_col = "銷貨金額" if "銷貨金額" in orders.columns else None
    if not amt_col:
        return None

    po_col = COL_GLOCOM_PO if COL_GLOCOM_PO in orders.columns else None
    factory_col = "工廠" if "工廠" in orders.columns else None
    date_col = "工廠下單日期" if "工廠下單日期" in orders.columns else None

    pn_target = str(part_number).strip()
    matches = orders[orders[pn_col].astype(str).str.strip() == pn_target].copy()
    if matches.empty:
        return None

    # 算每筆的匹配分數
    def _row_score(row):
        score = 0
        if factory_col and factory_short:
            rec_factory = str(row.get(factory_col, "") or "").strip()
            if rec_factory == factory_short:
                score += 5
        if po_col and po_prefix:
            rec_po = str(row.get(po_col, "") or "").strip()
            if _po_prefix(rec_po) == po_prefix:
                score += 10
        return score

    matches["_score"] = matches.apply(_row_score, axis=1)
    if date_col and date_col in matches.columns:
        matches["_date"] = pd.to_datetime(matches[date_col], errors="coerce")
        matches = matches.sort_values(["_score", "_date"], ascending=[False, False])
    else:
        matches = matches.sort_values("_score", ascending=False)

    # 從最佳匹配往下找,第一個有效的(qty>0 且 amt>0)就回傳
    for _, row in matches.iterrows():
        try:
            qty_raw = row.get(qty_col, 0)
            amt_raw = row.get(amt_col, 0)
            qty = float(qty_raw) if qty_raw not in (None, "") else 0.0
            amt = float(amt_raw) if amt_raw not in (None, "") else 0.0
            if qty > 0 and amt > 0:
                return round(amt / qty, 4)
        except (ValueError, TypeError):
            continue
    return None


def create_teable_records(table_url: str, headers: dict, fields_list: list) -> dict:
    result = {"success": 0, "failed": 0, "errors": [], "created_ids": []}
    if not fields_list:
        return result

    payload = {
        "fieldKeyType": "name",
        "typecast": True,
        "records": [{"fields": f} for f in fields_list],
    }

    try:
        r = requests.post(table_url, headers=headers, json=payload, timeout=30)
        if r.status_code in (200, 201):
            data = r.json()
            records = data.get("records", []) or []
            result["success"] = len(records) if records else len(fields_list)
            result["created_ids"] = [rec.get("id", "") for rec in records]
        else:
            result["failed"] = len(fields_list)
            result["errors"].append(f"HTTP {r.status_code}: {r.text[:500]}")
    except Exception as e:
        result["failed"] = len(fields_list)
        result["errors"].append(f"Exception: {e}")

    return result


# ─── Teable 查同料號歷史 ─────────────────────────
def load_spec_history() -> dict:
    """
    讀 data/spec_history.json (從歷史 RTF/DOCX 訂單檔抽出來的注意事項對應表)。

    ★ v3.9.8: 載入時統一 normalize 所有 spec_text(去 tab + 多餘空白),
              避免歷史資料的 "1.\\t..." 對齊 tab 印到 PDF 變成大空白。
    """
    if not SPEC_HISTORY_JSON.exists():
        return {}
    try:
        with open(SPEC_HISTORY_JSON, "r", encoding="utf-8") as f:
            data = json.load(f)
        spec_history = data.get("spec_history", {}) or {}

        # ★ v3.9.8: 整個 dict 走一遍,清掉每筆 spec_text 的 tab
        for pn, pn_data in spec_history.items():
            if not isinstance(pn_data, dict):
                continue
            latest = pn_data.get("latest")
            if isinstance(latest, dict) and "spec_text" in latest:
                latest["spec_text"] = _normalize_spec_text(latest["spec_text"])
            history_list = pn_data.get("history")
            if isinstance(history_list, list):
                for record in history_list:
                    if isinstance(record, dict) and "spec_text" in record:
                        record["spec_text"] = _normalize_spec_text(record["spec_text"])

        return spec_history
    except Exception as e:
        st.warning(f"spec_history.json 讀取失敗:{e}")
        return {}


def fetch_previous_spec(orders: pd.DataFrame, part_number: str, factory_short: str = ""):
    """
    查同 P/N 最近一筆訂單的規格。
    
    優先順序:
    1. data/spec_history.json (從歷史 RTF/DOCX 解析)
    2. Teable 主表「客戶要求注意事項」欄
    
    同工廠優先(如果有指定 factory_short)。

    ★ v3.9.8: 從 Teable 抓的 spec 也統一 normalize(雖然主表通常是 Sandy 手打的乾淨內容,
              但保險起見一致處理)。
    """
    results = []
    pn_target = str(part_number).strip()
    if not pn_target:
        return results
    
    # ─── 1. 從 spec_history.json 查 ─
    spec_history = load_spec_history()
    if pn_target in spec_history:
        history_records = spec_history[pn_target].get("history", [])
        # 同工廠優先
        if factory_short:
            same_factory = [r for r in history_records if r.get("factory") == factory_short]
            other = [r for r in history_records if r.get("factory") != factory_short]
            ordered = same_factory + other
        else:
            ordered = history_records
        for r in ordered:
            results.append({
                "po_no": r.get("po", ""),
                "spec": r.get("spec_text", ""),  # 已在 load_spec_history 時 normalize
                "factory": r.get("factory", ""),
                "date": r.get("date", ""),
                "source": "history_file",
            })
    
    # ─── 2. 若 spec_history 沒命中,再查 Teable 主表 ─
    if not results and not orders.empty:
        pn_col = "P/N"
        spec_col = "客戶要求注意事項"
        po_col = COL_GLOCOM_PO
        date_col = "工廠下單日期"
        factory_col = "工廠"
        
        if pn_col in orders.columns:
            matches = orders[orders[pn_col].astype(str).str.strip() == pn_target].copy()
            if not matches.empty:
                if date_col in matches.columns:
                    matches[date_col + "_dt"] = pd.to_datetime(matches[date_col], errors="coerce")
                    matches = matches.sort_values(date_col + "_dt", ascending=False)
                
                seen_pos = set()
                for _, row in matches.iterrows():
                    po = str(row.get(po_col, "")).strip() if po_col in row.index else ""
                    if not po or po in seen_pos:
                        continue
                    seen_pos.add(po)
                    spec = str(row.get(spec_col, "")).strip() if spec_col in row.index else ""
                    if pd.isna(spec) or spec.lower() == "nan":
                        spec = ""
                    spec = _normalize_spec_text(spec)  # ★ v3.9.8
                    factory = str(row.get(factory_col, "")).strip() if factory_col in row.index else ""
                    date_str = str(row.get(date_col, "")).strip() if date_col in row.index else ""
                    results.append({
                        "po_no": po,
                        "spec": spec,
                        "factory": factory,
                        "date": date_str,
                        "source": "teable",
                    })
                    if len(results) >= 5:
                        break
    
    return results


# ─── 規格組字串 ──────────────────────────────────
def build_full_spec_oneline(
    selections: dict,
    extra_text: str,
    engineering_confirm: str = "Working Gerber 承認",  # ★ v3.9
    include_footer: bool = True,                       # ★ v3.9
) -> str:
    """
    組合新料號規格(3 段式):
    [第1段] 工程確認語(預設 Working Gerber 承認)
    [第2段] 7 選單詳細規格 + 其他補充
    [第3段] 結尾標準語(預設加)
    """
    sections = []
    
    # 第 1 段:工程確認語
    confirm_text = ENGINEERING_CONFIRM_OPTIONS.get(engineering_confirm, "").strip()
    if confirm_text:
        sections.append(confirm_text)
    
    # 第 2 段:詳細規格 + 補充
    parts = []
    for cat, val in selections.items():
        if val:
            if cat in ("Tg",):
                parts.append(val)
            else:
                parts.append(f"{cat}: {val}")
    if extra_text and extra_text.strip():
        parts.append(extra_text.strip())
    
    if parts:
        sections.append("; ".join(parts))
    
    # 第 3 段:結尾標準語
    if include_footer:
        sections.append(FACTORY_PO_FOOTER)
    
    # 3 段之間用換行分隔(可讀性更好)
    return "\n".join(sections)


def build_po_context_from_new_order(
    new_po_no, parsed, factory_data, factory_short, issuing_company,
    factory_unit_prices, factory_due_dates, item_specs,
    purchase_responsible, order_date, is_revised,
    nre_settings=None,  # ★ v3.8 新增
):
    """
    nre_settings (dict, optional):
        {
            "has_nre": bool,
            "nre_amount": float (NT$),
            "nre_target_pn": str,
            "detected_nre_pns": list of str,  # 客戶 PDF 上的 NRE 品項料號(這些不該印工廠 PDF)
        }
    """
    nre_settings = nre_settings or {}
    nre_mode = nre_settings.get("mode", "merge")
    has_nre = nre_settings.get("has_nre", False)
    nre_amount = float(nre_settings.get("nre_amount", 0) or 0)
    nre_target_pn = nre_settings.get("nre_target_pn")
    detected_nre_pns = set(nre_settings.get("detected_nre_pns", []))
    factory_currency = nre_settings.get("factory_currency", "NTD")
    # ★ v3.9.5: 主品項合併 NRE 前的原始工廠單價(顯示在 PDF 註記裡)
    original_unit_price = float(nre_settings.get("original_unit_price", 0.0) or 0.0)
    
    apply_merge = (nre_mode == "merge" and has_nre)
    skip_nre_in_pdf = (
        (nre_mode == "merge") or
        (nre_mode == "separate" and not has_nre)
    )
    
    items = []
    for it in parsed.items:
        # NRE 品項是否跳過
        if it.part_number in detected_nre_pns and skip_nre_in_pdf:
            continue
        
        f_price = float(factory_unit_prices.get(it.part_number, 0.0))
        f_due = factory_due_dates.get(it.part_number)
        spec_text = (item_specs.get(it.part_number, "") or "").strip()

        # ★ v3.9.9: 保險再 normalize 一次(去 tab + 多餘空白)
        spec_text = _normalize_spec_text(spec_text)
        # ★ v3.9.9: 編號項目前加 ZWSP,避免 Word/LibreOffice list auto-format
        spec_text = _add_zwsp_to_numbered_lines(spec_text)

        # spec_text 寫進 docx 前先 escape & < >
        spec_text_for_docx = escape_for_docx(spec_text)
        
        # ★ v3.9.5: NRE 註記改新格式 — 含「原單價 + NRE 攤提」完整計算過程
        # 範例:
        #   (內含
        #   單價
        #   NTD102/pcs
        #   + NRE:
        #   NTD14,000;
        #   NTD28/pcs)
        unit_price_note = ""
        if apply_merge and it.part_number == nre_target_pn and nre_amount > 0 and it.quantity > 0:
            nre_per_pcs = nre_amount / it.quantity
            cur = factory_currency  # 已經是 NTD 或 USD
            unit_price_note = (
                f"\n(內含\n"
                f"單價\n"
                f"{cur}{original_unit_price:.0f}/pcs\n"
                f"+ NRE:\n"
                f"{cur}{nre_amount:,.0f};\n"
                f"{cur}{nre_per_pcs:.0f}/pcs)"
            )
        unit_price_note_for_docx = escape_for_docx(unit_price_note)
        
        items.append({
            "part_number": it.part_number,
            "spec_text": spec_text_for_docx,
            "quantity": it.quantity,
            "panel_qty": None,
            "unit_price": f_price,
            "unit_price_note": unit_price_note_for_docx,
            "amount": round(f_price * it.quantity, 2),
            "delivery_date": f_due,
            "delivery_note": "",
        })

    return {
        "po_no": new_po_no,
        "order_type": parse_glocom_po_no(new_po_no).get("order_type", ""),
        "issuing_company": issuing_company,
        "order_date": order_date,
        "customer_name": parsed.customer_name,
        "customer_po_no": parsed.customer_po_no,
        "factory_short": factory_short,
        "factory": factory_data,
        "items": items,
        "total_amount": sum(it["amount"] for it in items),
        "currency": factory_data.get("default_currency", "NT$"),
        "payment_terms": factory_data.get("default_payment_terms", ""),
        "shipment_method": factory_data.get("default_shipment", "待通知"),
        "ship_to_default": factory_data.get("default_ship_to", "待通知"),
        "purchase_responsible": purchase_responsible,
        "is_revised": is_revised,
    }


def _reset_form():
    keys_to_clear = [k for k in list(st.session_state.keys())
                     if k.startswith("fpo_") or k == "customer_po_uploader"]
    for k in keys_to_clear:
        try:
            del st.session_state[k]
        except KeyError:
            pass


# ─── Step 5 規格輸入(每品項)──────────────────
def render_spec_input_for_item(
    idx: int,
    item,
    orders: pd.DataFrame,
    current_po_prefix: str = "",  # ★ v3.9.11: GC / ET / EW
) -> str:
    final_key = f"fpo_spec_final_{idx}"
    if final_key not in st.session_state:
        st.session_state[final_key] = ""

    type_key = f"fpo_spec_type_{idx}"
    if type_key not in st.session_state:
        st.session_state[type_key] = "舊料號"

    applied_old_default_key = f"fpo_spec_old_applied_{idx}"
    # ★ v3.9.8: 保存「歷史規格 default 原版」,供「規格變更」按鈕重組用
    history_default_key = f"_fpo_spec_history_default_{idx}"

    with st.container(border=True):
        st.markdown(f"**品項 {idx + 1}:`{item.part_number}`**(數量 {item.quantity:,})")

        cols_type = st.columns([2, 5])
        with cols_type[0]:
            new_type = st.radio(
                "料號類型", ["舊料號", "新料號"],
                key=type_key, horizontal=True,
            )

        if new_type == "舊料號":
            with cols_type[1]:
                st.caption("👉 系統會從歷史訂單檔(模糊比對同料號)智能分析,自動帶入合併版規格。請檢查紅色疑點。")

            # ★ v3.9.10: needs_apply 條件 — 比舊版的單一旗標更 robust
            #   1. 從來沒套過(applied_old_default_key=False)→ 套
            #   2. final_key 為空(切過新料號清空後切回來)→ 重套
            #   3. Step 3 工廠變了,且 Sandy 還沒手動編輯過 final_key → 重套(切到同工廠版本)
            # ★ v3.9.11 加上:字首(GC/ET/EW)變了 → 也觸發重套
            current_factory_step3 = st.session_state.get("fpo_factory_short", "")
            already_applied = st.session_state.get(applied_old_default_key, False)
            current_final_value = (st.session_state.get(final_key, "") or "").strip()
            last_history_default = (st.session_state.get(history_default_key, "") or "").strip()
            last_factory_for_default_key = f"_fpo_spec_last_factory_{idx}"
            last_factory_for_default = st.session_state.get(last_factory_for_default_key, "__INITIAL__")
            last_prefix_for_default_key = f"_fpo_spec_last_prefix_{idx}"  # ★ v3.9.11
            last_prefix_for_default = st.session_state.get(last_prefix_for_default_key, "__INITIAL__")

            factory_changed = (
                last_factory_for_default != "__INITIAL__"
                and last_factory_for_default != current_factory_step3
            )
            prefix_changed = (  # ★ v3.9.11
                last_prefix_for_default != "__INITIAL__"
                and last_prefix_for_default != current_po_prefix
            )
            final_unchanged_by_user = (current_final_value == last_history_default)

            needs_apply = (
                (not already_applied)                                       # 從來沒套過
                or (not current_final_value)                                # final 為空 → 重套
                or ((factory_changed or prefix_changed) and final_unchanged_by_user)  # 工廠/字首變且沒被改過
            )

            if needs_apply:
                current_factory = current_factory_step3
                
                # ─── ★ v3.7: 智能歷史分析(模糊料號 + 行頻率分析) ─
                from core.spec_intelligence import (
                    analyze_spec_history,
                    build_smart_spec,
                    find_similar_part_numbers,
                )
                
                spec_history_dict = load_spec_history()
                
                # 找模糊料號(包含原始料號自己)
                target_pn = item.part_number
                matched_pn = None
                history_records = []
                
                # 1. 直接 hit
                if target_pn in spec_history_dict:
                    matched_pn = target_pn
                    history_records = spec_history_dict[target_pn].get("history", [])
                else:
                    # 2. 模糊比對
                    similar = find_similar_part_numbers(
                        target_pn,
                        list(spec_history_dict.keys()),
                        strict=False,  # 寬鬆比對
                    )
                    if similar:
                        # 取相似料號中歷史最多的當主要參考
                        similar.sort(key=lambda p: -len(spec_history_dict[p].get("history", [])))
                        matched_pn = similar[0]
                        history_records = spec_history_dict[matched_pn].get("history", [])
                
                # 3. 智能分析
                if history_records:
                    analysis = analyze_spec_history(target_pn, history_records)
                    smart_spec = build_smart_spec(analysis)
                    # ★ v3.9.8: smart_spec 也保險再 normalize 一次
                    smart_spec = _normalize_spec_text(smart_spec)

                    # ★ v3.9.11: 雙條件排序 — 同字首+同工廠 > 同字首 > 同工廠 > 其他
                    #   Sandy 切 GC/ET/EW 字首 + 切工廠時都會 follow
                    def _match_score(record):
                        score = 0
                        rec_factory = (record.get("factory") or "").strip()
                        rec_prefix = _po_prefix(record.get("po", ""))
                        if current_po_prefix and rec_prefix == current_po_prefix:
                            score += 10  # 同字首加 10
                        if current_factory and rec_factory == current_factory:
                            score += 5   # 同工廠加 5
                        return (score, record.get("date") or "")

                    sorted_recs = sorted(history_records, key=_match_score, reverse=True)
                    top_score = _match_score(sorted_recs[0])[0] if sorted_recs else 0
                    if sorted_recs and top_score > 0:
                        # 至少有同字首 OR 同工廠 → 用該筆
                        chosen = sorted_recs[0]
                        chosen_spec = _normalize_spec_text(
                            (chosen.get("spec_text") or "").strip()
                        )
                        # ★ v3.9.12: 不加重複前綴 — 歷史 spec_text 第一行通常就是「舊料號;...」開頭
                        default_text = _ensure_starts_with_jiu_liao(chosen_spec)
                    else:
                        # 沒任何匹配 → smart_spec 合併版
                        # ★ v3.9.12: 同樣處理
                        default_text = _ensure_starts_with_jiu_liao(smart_spec)
                    
                    # ★ v3.9.9: 同時把所有歷史工廠 + 對應規格存起來,供「歷史工廠」UI 使用
                    factory_choices = []
                    seen_keys = set()
                    for r in history_records:
                        factory = (r.get("factory") or "").strip()
                        po = (r.get("po") or "").strip()
                        date_str = (r.get("date") or "").strip()
                        spec = _normalize_spec_text((r.get("spec_text") or "").strip())
                        # 同 (工廠+規格) 視為同一筆,避免重複
                        dedup_key = (factory, spec)
                        if dedup_key in seen_keys:
                            continue
                        seen_keys.add(dedup_key)
                        factory_choices.append({
                            "factory": factory or "(未填工廠)",
                            "po": po,
                            "date": date_str,
                            "spec_text": spec,
                            "po_prefix": _po_prefix(po),  # ★ v3.9.11
                        })
                    st.session_state[f"_fpo_factory_choices_{idx}"] = factory_choices

                    # 存分析結果到 session_state(下面 UI 會顯示)
                    st.session_state[f"_fpo_v37_analysis_{idx}"] = {
                        "matched_pn": matched_pn,
                        "is_fuzzy": (matched_pn != target_pn),
                        "total_history": analysis.total_history,
                        "core_count": len(analysis.core_lines),
                        "new_count": len(analysis.new_lines),
                        "dropped_count": len(analysis.dropped_lines),
                        "has_warnings": analysis.has_warnings,
                        "latest_po": analysis.latest_record.get("po", ""),
                        "latest_factory": analysis.latest_record.get("factory", ""),
                        "latest_date": analysis.latest_record.get("date", ""),
                        "annotated_lines": [
                            {
                                "text": al.text,
                                "category": al.category,
                                "explanation": al.explanation,
                                "color_code": al.color_code,
                            }
                            for al in analysis.annotated_lines
                        ],
                    }
                else:
                    # 沒找到任何歷史 — fallback 到原本邏輯查 Teable
                    hist_hits = fetch_previous_spec(orders, item.part_number, factory_short=current_factory)
                    if hist_hits and hist_hits[0].get("spec"):
                        # ★ v3.9.12: 不加重複前綴
                        default_text = _ensure_starts_with_jiu_liao(hist_hits[0]['spec'].strip())
                        # ★ v3.9.9: Teable fallback 也建 factory_choices
                        factory_choices = []
                        seen_keys = set()
                        for h in hist_hits:
                            factory = (h.get("factory") or "").strip()
                            spec = (h.get("spec") or "").strip()
                            dedup_key = (factory, spec)
                            if dedup_key in seen_keys or not spec:
                                continue
                            seen_keys.add(dedup_key)
                            factory_choices.append({
                                "factory": factory or "(未填工廠)",
                                "po": h.get("po_no", ""),
                                "date": h.get("date", ""),
                                "spec_text": spec,
                                "po_prefix": _po_prefix(h.get("po_no", "")),  # ★ v3.9.11
                            })
                        st.session_state[f"_fpo_factory_choices_{idx}"] = factory_choices

                        st.session_state[f"_fpo_v37_analysis_{idx}"] = {
                            "matched_pn": item.part_number,
                            "is_fuzzy": False,
                            "total_history": 1,
                            "core_count": 0, "new_count": 0, "dropped_count": 0,
                            "has_warnings": False,
                            "latest_po": hist_hits[0].get("po_no", ""),
                            "latest_factory": hist_hits[0].get("factory", ""),
                            "latest_date": hist_hits[0].get("date", ""),
                            "annotated_lines": [],
                            "from_teable": True,
                        }
                    else:
                        default_text = "舊料號"
                        st.session_state[f"_fpo_v37_analysis_{idx}"] = None
                        st.session_state[f"_fpo_factory_choices_{idx}"] = []  # ★ v3.9.9
                
                st.session_state[final_key] = default_text
                # ★ v3.9.8: 同時保存 history_default 原版,供規格變更按鈕重組用
                st.session_state[history_default_key] = default_text
                st.session_state[applied_old_default_key] = True
                # ★ v3.9.10: 記住套用 default 時的工廠,下次切工廠時偵測差異
                st.session_state[last_factory_for_default_key] = current_factory_step3
                # ★ v3.9.11: 記住套用 default 時的字首,下次切字首時偵測差異
                st.session_state[last_prefix_for_default_key] = current_po_prefix
            
            # ─── 顯示 v3.7 分析結果 ─
            v37 = st.session_state.get(f"_fpo_v37_analysis_{idx}")
            if v37:
                # 來源標示
                if v37.get("from_teable"):
                    src_label = "Teable 主表"
                elif v37.get("is_fuzzy"):
                    src_label = f"歷史訂單檔(模糊比對到 `{v37['matched_pn']}`)"
                else:
                    src_label = "歷史訂單檔"
                
                src_info = f"**{v37['latest_po']}**"
                if v37.get("latest_factory"):
                    src_info += f" / {v37['latest_factory']}"
                if v37.get("latest_date"):
                    src_info += f" / {v37['latest_date']}"

                # ★ v3.9.14: 用 success 框讓來源資訊更顯眼
                st.success(
                    f"📌 **此規格 default 來源**:{src_info}({src_label})  \n"
                    f"💡 切 Step 3 工廠/字首會自動用「同字首+同工廠優先」的雙條件重新抓。"
                )
                
                # 統計徽章
                total_h = v37.get("total_history", 0)
                if total_h >= 1:
                    badge_cols = st.columns(4)
                    badge_cols[0].metric("📚 歷史筆數", total_h)
                    badge_cols[1].metric("🟢 核心要求", v37.get("core_count", 0))
                    badge_cols[2].metric("🔵 最新加的", v37.get("new_count", 0))
                    badge_cols[3].metric(
                        "🔴 疑點",
                        v37.get("dropped_count", 0),
                        delta=("⚠️ 請檢查" if v37.get("has_warnings") else None),
                        delta_color="inverse" if v37.get("has_warnings") else "off",
                    )
                
                # 警告區塊(若有 DROPPED)
                if v37.get("has_warnings"):
                    dropped_lines = [
                        a for a in v37.get("annotated_lines", [])
                        if a.get("category") == "DROPPED"
                    ]
                    with st.expander(
                        f"🔴 ⚠️ {len(dropped_lines)} 條規格在舊版有但最新沒寫(請確認)",
                        expanded=True,
                    ):
                        st.warning(
                            "以下規格出現在較舊的訂單,但最新一筆沒寫。"
                            "可能是:(a) 客戶要求變更了 / (b) 最新單漏寫。"
                            "**請確認哪些要保留**,規格框已自動加入並標 `?`。"
                        )
                        for a in dropped_lines:
                            st.markdown(f"- `{a['text'].strip()}` — {a['explanation']}")
                
                # 詳細分析(可展開)
                if v37.get("annotated_lines"):
                    with st.expander(f"📊 詳細規格分析(展開看每行的歷史身份)", expanded=False):
                        for a in v37.get("annotated_lines", []):
                            st.markdown(
                                f"{a['color_code']} `{a['text'].strip()[:80]}`  \n"
                                f"&nbsp;&nbsp;&nbsp;&nbsp;_{a['explanation']}_"
                            )
            else:
                st.caption("⚠️ 沒找到同料號的歷史訂單,請手動填注意事項。")

            # ─── ★ v3.9.8: 規格變更欄位 ─────────────
            st.divider()

            # ─── ★ v3.9.9: 歷史工廠 expander ─────────
            factory_choices = st.session_state.get(f"_fpo_factory_choices_{idx}", []) or []
            if factory_choices:
                current_factory_step3 = st.session_state.get("fpo_factory_short", "")
                # 同工廠優先 — 排在最前面
                same_factory = [c for c in factory_choices if c.get("factory") == current_factory_step3]
                other = [c for c in factory_choices if c.get("factory") != current_factory_step3]
                ordered_choices = same_factory + other

                expander_label = (
                    f"🏭 歷史工廠({len(ordered_choices)} 筆,可切換看不同工廠的規格,並一鍵套用)"
                )
                # ★ v3.9.10: 永遠預設展開,讓 Sandy 一進 Step 5 就看到歷史工廠
                with st.expander(expander_label, expanded=True):
                    st.caption(
                        "💡 切 Step 3 工廠時,如果歷史有同工廠紀錄,系統會自動帶該工廠的最近一筆。"
                        "想看不同工廠的版本就點下面對應的「⬇ 套用此筆」。"
                    )
                    for c_idx, choice in enumerate(ordered_choices):
                        is_same = (choice.get("factory") == current_factory_step3 and current_factory_step3)
                        badge = "🟢 同工廠 " if is_same else "⚪ "
                        # ★ v3.9.11: 字首 badge
                        prefix_emoji = {
                            "GC": "🟦 GC",
                            "ET": "🟩 ET",
                            "EW": "🟧 EW",
                        }.get(choice.get("po_prefix", ""), "⬛ ?")
                        # 如果這筆字首跟 Step 3 選的字首一樣,加 ★ 標記
                        is_same_prefix = (choice.get("po_prefix", "") == current_po_prefix and current_po_prefix)
                        prefix_star = " ★" if is_same_prefix else ""

                        sub_cols = st.columns([4, 1])
                        with sub_cols[0]:
                            st.markdown(
                                f"{badge}{prefix_emoji}{prefix_star} | "
                                f"**{choice['factory']}** | "
                                f"`{choice['po']}` | {choice.get('date', '') or '(無日期)'}"
                            )
                            if choice.get("spec_text"):
                                st.code(choice["spec_text"], language=None)
                            else:
                                st.caption("(此筆沒填規格)")
                        with sub_cols[1]:
                            st.write("")  # 對齊
                            if st.button(
                                "⬇ 套用此筆",
                                key=f"fpo_apply_factory_{idx}_{c_idx}",
                                use_container_width=True,
                                disabled=not choice.get("spec_text"),
                            ):
                                new_default = f"舊料號\n\n注意:\n{choice['spec_text'].strip()}"
                                st.session_state[final_key] = new_default
                                # 同步更新 history_default,讓「規格變更」按鈕從這個新 base 重組
                                st.session_state[history_default_key] = new_default
                                # 清掉「規格變更」欄位內容(避免殘留上一次的變更)
                                st.session_state[f"fpo_spec_change_{idx}"] = ""
                                st.rerun()
                        st.divider()

            # ─── 📝 這次規格變更 ──────────────────────
            st.markdown("##### 📝 這次規格變更(選填)")
            st.caption(
                "👉 用在「舊料號但這次規格有變更」的情況,譬如 "
                '`表面處理改 ENIG 2u"` / `改為 4 up array` / `板厚改 1.6mm`。  \n'
                "**打完按 Enter 就會自動套用**(也可按右邊「⬇ 套用變更」)。  \n"
                "寫回 Teable 後也會更新 `spec_history.json`,下次同料號自動帶新版規格。"
            )
            change_key = f"fpo_spec_change_{idx}"
            cols_chg = st.columns([4, 1])
            with cols_chg[0]:
                spec_change_text = st.text_input(
                    "規格變更",
                    key=change_key,
                    placeholder='例: 表面處理改 ENIG 2u"',
                    label_visibility="collapsed",
                    # ★ v3.9.10: 按 Enter 自動觸發 callback,不用再點按鈕
                    on_change=_on_spec_change_apply,
                    args=(history_default_key, change_key, final_key),
                )
            with cols_chg[1]:
                if st.button(
                    "⬇ 套用變更",
                    key=f"fpo_apply_change_{idx}",
                    use_container_width=True,
                    type="primary",
                ):
                    # 跟 callback 同樣的邏輯(備援用,有些使用者習慣點按鈕)
                    _on_spec_change_apply(history_default_key, change_key, final_key)
                    st.rerun()

        else:  # 新料號
            with cols_type[1]:
                st.caption("👉 新料號:用下方 3 個工具填完整規格,或直接在最終規格框打字。")
            st.session_state[applied_old_default_key] = False
            
            # ★ v3.9.1: 切換到新料號時,如果 final_key 還停留在舊料號的內容(以「舊料號」開頭),
            #            清空讓 Sandy 從頭開始
            current_final = st.session_state.get(final_key, "")
            last_type_key = f"_fpo_spec_last_type_{idx}"
            last_type = st.session_state.get(last_type_key, new_type)
            
            if last_type != new_type:
                # 類型有變(舊→新 或 新→舊),清掉舊內容
                st.session_state[final_key] = ""
                current_final = ""
            elif current_final.strip().startswith("舊料號") and not current_final.strip() == "":
                # 同一輪剛從舊料號切過來,殘留「舊料號」相關文字 → 清掉
                if current_final.strip() in ("舊料號", "舊料號\n", "舊料號\n\n注意:"):
                    st.session_state[final_key] = ""
            
            # 紀錄目前類型
            st.session_state[last_type_key] = new_type

        if new_type == "新料號":
            # ★ v3.9.2: 7 選單預設展開(不再藏在 tab),
            #            查詢歷史 + 多列貼上 改放底下 expander
            
            # ─── 第 1 段:工程確認語 ─
            st.markdown("**第 1 段:工程確認語**")
            eng_confirm_key = f"fpo_spec_eng_confirm_{idx}"
            eng_confirm = st.radio(
                "工程確認語",
                options=list(ENGINEERING_CONFIRM_OPTIONS.keys()),
                index=0,
                horizontal=True,
                key=eng_confirm_key,
                label_visibility="collapsed",
            )
            preview_text = ENGINEERING_CONFIRM_OPTIONS.get(eng_confirm, "")
            if preview_text:
                st.caption(f"預覽: _{preview_text}_")
            else:
                st.caption("(此段不寫)")
            
            st.divider()
            
            # ─── 第 2 段:7 選單 + 補充 ─
            st.markdown("**第 2 段:詳細規格**")
            cols_a = st.columns(4)
            cols_b = st.columns(4)
            selections = {}
            cats = list(SPEC_OPTIONS_FULL.keys())
            for i, cat in enumerate(cats):
                target_col = cols_a[i] if i < 4 else cols_b[i - 4]
                with target_col:
                    selections[cat] = st.selectbox(
                        cat, [""] + SPEC_OPTIONS_FULL[cat],
                        key=f"fpo_spec_full_{idx}_{cat}",
                    )

            extra = st.text_input(
                "其他規格 / 補充",
                placeholder="例如: 排版 4up panel",
                key=f"fpo_spec_full_extra_{idx}",
            )
            
            st.divider()
            
            # ─── 第 3 段:結尾標準語(強制加) ─
            st.markdown("**第 3 段:結尾標準語(必加)**")
            st.caption(f"_{FACTORY_PO_FOOTER}_")

            # ─── 組合 + 套用 ─
            st.divider()
            preview = build_full_spec_oneline(
                selections, extra,
                engineering_confirm=eng_confirm,
                include_footer=True,
            )
            cols_btn = st.columns([3, 1])
            with cols_btn[0]:
                if preview:
                    st.markdown(f"📝 **完整規格預覽**:")
                    st.code(preview, language=None)
            with cols_btn[1]:
                if st.button(
                    "⬇ 套用到下方規格框",
                    key=f"fpo_apply_full_{idx}",
                    disabled=not preview,
                    use_container_width=True,
                    type="primary",
                ):
                    st.session_state[final_key] = preview
                    st.rerun()
            
            # ─── 進階工具(預設摺疊) ─
            with st.expander("🔧 進階工具:從歷史訂單帶入 / 從 ERP 多列貼上", expanded=False):
                st.markdown("##### ⚡ 從歷史訂單帶入規格")
                st.caption(f"從 Teable 主表 / 歷史訂單檔查 P/N『{item.part_number}』")
                if st.button(f"🔍 查詢歷史", key=f"fpo_query_btn_{idx}"):
                    hits = fetch_previous_spec(orders, item.part_number)
                    st.session_state[f"fpo_query_hits_{idx}"] = hits

                hits = st.session_state.get(f"fpo_query_hits_{idx}")
                if hits is not None:
                    if not hits:
                        st.warning(f"沒找到 P/N『{item.part_number}』的歷史訂單")
                    else:
                        st.info(f"找到 **{len(hits)}** 筆歷史訂單(由新到舊)")
                        for h_idx, h in enumerate(hits):
                            with st.expander(
                                f"📦 {h['po_no']}  |  {h['factory']}  |  {h['date']}",
                                expanded=(h_idx == 0),
                            ):
                                if h["spec"]:
                                    st.code(h["spec"], language=None)
                                    if st.button(
                                        f"⬇ 套用此規格",
                                        key=f"fpo_apply_hist_{idx}_{h_idx}",
                                        type="primary",
                                    ):
                                        st.session_state[final_key] = h["spec"]
                                        st.rerun()
                                else:
                                    st.caption("(此筆訂單沒填規格)")
                
                st.divider()
                
                st.markdown("##### 📋 從 ERP 多列複製貼上")
                st.caption("最多 8 列,按下「組合」會用『; 』串成一列")
                paste_lines = []
                cols_paste = st.columns(2)
                for line_idx in range(8):
                    target = cols_paste[line_idx % 2]
                    with target:
                        line = st.text_input(
                            f"第 {line_idx + 1} 列",
                            key=f"fpo_paste_{idx}_{line_idx}",
                            placeholder="例如: 8L, FR4, 1.6mm, 1up",
                        )
                    if line and line.strip():
                        paste_lines.append(line.strip())

                cols_btn_b = st.columns([3, 1])
                with cols_btn_b[0]:
                    if paste_lines:
                        preview_paste = "; ".join(paste_lines)
                        st.caption(f"預覽: `{preview_paste}`")
                with cols_btn_b[1]:
                    if st.button(
                        "⬇ 組合並套用",
                        key=f"fpo_combine_paste_{idx}",
                        disabled=not paste_lines,
                        use_container_width=True,
                        type="primary",
                    ):
                        st.session_state[final_key] = "; ".join(paste_lines)
                        st.rerun()

        st.divider()
        st.markdown("##### 📋 最終規格(印在 PO 上,**可直接在框內打字編輯**)")
        if new_type == "舊料號":
            spec_placeholder = "舊料號\n注意:\n1. ...\n2. ..."
        else:
            spec_placeholder = (
                "上方 7 選單填好後按「⬇ 套用到下方規格框」會自動填入,可再修改。\n"
                "或直接在這裡打字。\n\n"
                "範例(自動產生):\n"
                "Working Gerber承認後,才可生產!...\n"
                "Material: 4L; Tg170; Board thickness: 1.6mm; ...\n"
                "須添加西拓UL logo & date code (YYWW);\n"
                "樣板: 除試錫板外,須另外提供樣板供備份."
            )

        st.text_area(
            "最終規格",
            key=final_key,
            label_visibility="collapsed",
            height=180,
            placeholder=spec_placeholder,
        )

    return st.session_state[final_key]


# ─── ★ v3.5: 即時撈 Teable 編號 + 計算下一編號 (有 cache) ─
def get_live_po_numbers(table_url: str, headers: dict, force_refresh: bool = False) -> set[str]:
    """
    即時撈 Teable 上所有訂單編號。30 秒 cache 避免每次 rerun 都打 API。
    
    Args:
        force_refresh: True 時強制重撈
    """
    cache_key = "_fpo_live_po_numbers"
    cache_time_key = "_fpo_live_po_numbers_t"
    
    now = datetime.now().timestamp()
    cached_time = st.session_state.get(cache_time_key, 0)
    cached_set = st.session_state.get(cache_key)
    
    if not force_refresh and cached_set is not None and (now - cached_time) < 30:
        return cached_set
    
    fresh_set = fetch_all_po_numbers_from_teable(table_url, headers)
    st.session_state[cache_key] = fresh_set
    st.session_state[cache_time_key] = now
    return fresh_set


# ─── 主入口 ───────────────────────────────────────
def render_factory_po_create_page(orders: pd.DataFrame, table_url: str, headers: dict):
    st.subheader("📝 建立工廠 PO(新流程)")
    st.caption("上傳客戶 PO PDF → 自動解析 → 選工廠/規格 → 產出工廠 PO PDF / 寫回 Teable")

    factories = load_factories()
    if not factories:
        st.error("data/factories.json 讀取失敗或無工廠資料")
        st.stop()

    # ─── Step 1 ─
    st.markdown("### Step 1. 上傳客戶 PO PDF")
    col_up, col_clear = st.columns([4, 1])
    with col_up:
        uploaded = st.file_uploader(
            "支援:WESCO / TIETO / GUDE / KCS / VORNE",
            type=["pdf"],
            key="customer_po_uploader",
        )
    with col_clear:
        st.write("")
        st.write("")
        if st.button("🗑️ 清除", use_container_width=True, key="fpo_clear_top"):
            _reset_form()
            st.rerun()

    if uploaded is None:
        st.info("👆 請上傳客戶 PO PDF。")
        st.stop()

    # ─── Step 2 ─
    try:
        with st.spinner("解析 PDF..."):
            pdf_bytes = uploaded.getvalue()
            parsed = parse_customer_po_from_pdf(pdf_bytes)
    except Exception as e:
        st.error(f"PDF 解析失敗:{e}")
        with st.expander("錯誤詳情"):
            st.code(traceback.format_exc())
        st.stop()

    st.markdown("### Step 2. 解析結果")
    parser_badge = {
        "WESCO": "🟢 WESCO 解析器", "TIETO": "🟢 TIETO 解析器",
        "GUDE": "🟢 GUDE 解析器", "KCS": "🟢 KCS 解析器",
        "VORNE": "🟢 VORNE 解析器", "GENERIC": "🟡 通用模式(欄位請手動補)",
    }
    st.caption(f"使用解析器:{parser_badge.get(parsed.parser_used, parsed.parser_used)}")
    if parsed.parse_warnings:
        for w in parsed.parse_warnings:
            st.warning(w)

    info_cols = st.columns(4)
    info_cols[0].metric("客戶", parsed.customer_name or "(未識別)")
    info_cols[1].metric("客戶 PO#", parsed.customer_po_no or "—")
    info_cols[2].metric("PO 日期", parsed.po_date or "—")
    info_cols[3].metric("總金額", f"{parsed.currency} {parsed.total_amount:,.2f}")

    if not parsed.items:
        st.error("沒有解析出品項,無法繼續。請檢查 PDF 是否正確。")
        with st.expander("📄 PDF 原始文字"):
            st.text(parsed.raw_text)
        st.stop()

    # 可編輯品項表
    st.markdown("#### 品項清單(可編輯)")
    items_df = pd.DataFrame([{
        "P/N": it.part_number, "描述": it.description, "數量": it.quantity,
        "客戶單價": it.unit_price, "客戶金額": it.amount,
        "客戶交期": it.delivery_date or "",
    } for it in parsed.items])
    edited_items_df = st.data_editor(
        items_df, use_container_width=True, hide_index=True,
        num_rows="fixed", key="fpo_items_editor",
    )
    for i, row in edited_items_df.iterrows():
        if i < len(parsed.items):
            parsed.items[i].part_number = str(row["P/N"]).strip()
            parsed.items[i].description = str(row["描述"]).strip()
            try: parsed.items[i].quantity = int(row["數量"])
            except: pass
            try: parsed.items[i].unit_price = float(row["客戶單價"])
            except: pass
            try: parsed.items[i].amount = float(row["客戶金額"])
            except: pass
            parsed.items[i].delivery_date = str(row["客戶交期"]).strip() or None

    # ─── ★ v3.9.17: 舊料號歷史工廠 — 自動套用 + 面板總是顯示 ─
    # ★ v3.9.19: 整段包 try/except,任何意外不讓 page 崩
    try:
        spec_history_data_for_panel = load_spec_history()
    from legacy_history import merge_legacy_into_spec_history
        spec_history_data = merge_legacy_into_spec_history(spec_history_data)
    except Exception as e:
        spec_history_data_for_panel = {}
        st.warning(f"⚠️ 載入 spec_history.json 失敗(已 fallback 空 dict):{e}")

    # 確保是 dict
    if not isinstance(spec_history_data_for_panel, dict):
        st.warning(
            f"⚠️ spec_history.json 不是預期的 dict 格式(實際:{type(spec_history_data_for_panel).__name__})"
        )
        spec_history_data_for_panel = {}

    try:
        pn_suggestions = _compute_old_pn_factory_suggestions(
            parsed.items, spec_history_data_for_panel
        )
        has_any_history = any(s.get("total", 0) > 0 for s in pn_suggestions.values())
        votes = _aggregate_factory_prefix_votes(pn_suggestions)
    except Exception as e:
        st.error(f"⚠️ 歷史工廠分析失敗:{e}")
        with st.expander("詳細錯誤"):
            st.code(traceback.format_exc())
        pn_suggestions = {}
        has_any_history = False
        votes = Counter()

    auto_applied_now = False
    auto_applied_combo = None
    auto_applied_count = 0

    try:
        # PDF identity sentinel — 同一份 PDF 只自動套一次
        auto_applied_key = "_fpo_auto_applied_factory_sentinel"
        pdf_sentinel = (
            f"{parsed.customer_po_no or ''}|{len(parsed.items)}|"
            f"{parsed.items[0].part_number if parsed.items else ''}"
        )

        if votes:
            best, best_count = votes.most_common(1)[0]
            best_factory, best_prefix = best
            last_applied_for = st.session_state.get(auto_applied_key)
            cur_factory = st.session_state.get("fpo_factory_short", "")
            cur_prefix = st.session_state.get("fpo_prefix", "")

            needs_auto = (
                last_applied_for != pdf_sentinel
                and best_factory in factories
                and (cur_factory != best_factory or cur_prefix != best_prefix)
            )
            if needs_auto:
                st.session_state["fpo_factory_short"] = best_factory
                st.session_state["fpo_prefix"] = best_prefix
                if best_prefix == "GC":
                    st.session_state["fpo_issuing"] = "GLOCOM"
                elif best_prefix in ("ET", "EW"):
                    st.session_state["fpo_issuing"] = "EUSWAY"
                st.session_state[auto_applied_key] = pdf_sentinel
                auto_applied_now = True
                auto_applied_combo = best
                auto_applied_count = best_count
    except Exception as e:
        st.warning(f"⚠️ 自動套用工廠失敗:{e}")

    # ─── 面板:永遠顯示(就算沒歷史也顯示診斷訊息)
    try:
        with st.container(border=True):
            st.markdown("##### 📋 舊料號歷史工廠")

            if not has_any_history:
                # 沒任何歷史 → 顯示診斷訊息讓 Sandy 知道為什麼沒自動切
                pn_list = ", ".join([
                    f"`{getattr(it, 'part_number', '')}`" for it in parsed.items
                    if getattr(it, "part_number", None)
                ])
                st.warning(
                    f"⚠️ 在 spec_history.json 裡找不到這些料號的歷史:{pn_list}  \n"
                    f"→ 系統無法自動建議工廠。請手動在下方 Step 3 選擇。"
                )
            else:
                if auto_applied_now and auto_applied_combo:
                    af, ap = auto_applied_combo
                    st.success(
                        f"✨ **已自動切到歷史最多的工廠 + 字首** —  \n"
                        f"工廠 **{af}** + 字首 **{ap}**(共 {auto_applied_count} 筆紀錄)"
                    )
                elif votes:
                    best, best_count = votes.most_common(1)[0]
                    af, ap = best
                    cur_f = st.session_state.get("fpo_factory_short", "")
                    cur_p = st.session_state.get("fpo_prefix", "")
                    if cur_f == af and cur_p == ap:
                        st.success(
                            f"✅ Step 3 已是多數歷史的選擇:**{af}** + **{ap}**"
                            f"(共 {best_count} 筆紀錄)"
                        )
                    else:
                        st.info(
                            f"💡 多數歷史建議:**{af}** + **{ap}**(共 {best_count} 筆紀錄)  \n"
                            f"目前 Step 3 是:**{cur_f}** + **{cur_p}**(Sandy 已手動切換,不再自動覆蓋)"
                        )

                # 細節表格:每個品項在 GC/ET/EW 三個字首下的歷史工廠分佈
                current_factory_for_panel = st.session_state.get("fpo_factory_short", "")
                current_prefix_for_panel = st.session_state.get("fpo_prefix", "")
                rows = []
                for pn, sug in pn_suggestions.items():
                    row = {"P/N": pn}
                    for prefix in ["GC", "ET", "EW"]:
                        emoji = {"GC": "🟦", "ET": "🟩", "EW": "🟧"}[prefix]
                        top_list = sug["by_prefix"][prefix].most_common(3)
                        if top_list:
                            parts = []
                            for f, c in top_list:
                                badge = " ★" if (
                                    f == current_factory_for_panel
                                    and prefix == current_prefix_for_panel
                                ) else ""
                                parts.append(f"{f}({c}){badge}")
                            row[f"{emoji} {prefix}"] = " · ".join(parts)
                        else:
                            row[f"{emoji} {prefix}"] = "—"
                    row["共"] = f"{sug['total']} 筆"
                    rows.append(row)
                st.dataframe(
                    pd.DataFrame(rows),
                    use_container_width=True,
                    hide_index=True,
                )

            # ★ v3.9.17: 診斷 expander — 顯示 spec_history.json 內容
            with st.expander("🔍 spec_history.json 診斷(展開看載入了什麼)", expanded=False):
                history_path = Path("data/spec_history.json")
                alt_paths = [
                    history_path,
                    Path("/mount/src/glocom-wip-v36/data/spec_history.json"),
                    Path(__file__).parent / "data" / "spec_history.json",
                ]
                actual_path = next((p for p in alt_paths if p.exists()), None)
                if actual_path:
                    file_size = actual_path.stat().st_size
                    st.caption(
                        f"檔案路徑:`{actual_path}` · 大小:{file_size:,} bytes · "
                        f"總料號數:{len(spec_history_data_for_panel)}"
                    )
                else:
                    st.error(
                        f"⚠️ 找不到 spec_history.json!試過的路徑:{', '.join(str(p) for p in alt_paths)}"
                    )

                for it in parsed.items:
                    pn = getattr(it, "part_number", None)
                    if not pn:
                        continue
                    records_raw = spec_history_data_for_panel.get(pn, [])
                    if isinstance(records_raw, dict):
                        records = [records_raw]
                    elif isinstance(records_raw, list):
                        records = records_raw
                    else:
                        records = []
                    records = [r for r in records if isinstance(r, dict)]
                    if records:
                        st.write(f"**`{pn}`** — 找到 **{len(records)}** 筆紀錄")
                        detail_rows = []
                        for r in records:
                            detail_rows.append({
                                "工廠": r.get("factory", ""),
                                "PO 編號": r.get("po", ""),
                                "字首": _po_prefix(r.get("po", "")),
                                "日期": r.get("date", ""),
                                "規格首行": (r.get("spec_text") or "").split("\n", 1)[0][:60],
                            })
                        st.dataframe(
                            pd.DataFrame(detail_rows),
                            use_container_width=True,
                            hide_index=True,
                        )
                    else:
                        raw_count = len(records_raw) if isinstance(records_raw, (list, dict)) else 0
                        if raw_count > 0:
                            st.write(
                                f"**`{pn}`** — 有 {raw_count} 筆原始紀錄但格式不是 dict,已跳過"
                            )
                            st.code(repr(records_raw)[:500], language="python")
                        else:
                            st.write(f"**`{pn}`** — ❌ 沒有任何歷史紀錄")
                            similar_keys = [
                                k for k in spec_history_data_for_panel.keys()
                                if pn[:6] in k or k[:6] in pn
                            ][:5]
                            if similar_keys:
                                st.caption(f"  spec_history 裡名稱接近的料號:{', '.join(similar_keys)}")
    except Exception as e:
        # ★ v3.9.19: 整個面板渲染掛掉也不能讓 page 崩
        st.error(f"⚠️ 歷史工廠面板渲染失敗(其他功能仍可使用):{e}")
        with st.expander("詳細錯誤"):
            st.code(traceback.format_exc())

    # ─── Step 3 ─
    st.markdown("### Step 3. 工廠 / 發出公司 / 字首 / 編號")
    s3_cols = st.columns(3)
    factory_keys = sorted(
        factories.keys(),
        key=lambda k: (factories[k].get("sort_priority", 999), k),
    )
    with s3_cols[0]:
        factory_short = st.selectbox(
            "工廠 *", factory_keys, key="fpo_factory_short",
        )
    factory_data = factories[factory_short]
    factory_region = factory_data.get("region", "Taiwan")

    missing_fields = []
    for k in ["factory_name", "address", "phone", "default_currency", "default_payment_terms"]:
        v = factory_data.get(k, "")
        if not v or "[請補]" in str(v):
            missing_fields.append(k)
    if missing_fields:
        st.warning(f"⚠️ 工廠「{factory_short}」資料不完整,缺:{', '.join(missing_fields)}")

    with s3_cols[1]:
        default_issuing = parsed.issuing_company_detected or "GLOCOM"
        issuing_options = ["GLOCOM", "EUSWAY"]
        issuing_idx = issuing_options.index(default_issuing) if default_issuing in issuing_options else 0
        issuing_company = st.selectbox(
            "發出公司 *", issuing_options, index=issuing_idx, key="fpo_issuing",
        )

    auto_internal_prefix = derive_prefix(issuing_company, factory_region)
    auto_display = display_prefix(auto_internal_prefix)
    prefix_options = ["GC", "ET", "EW"]
    with s3_cols[2]:
        prefix_idx = prefix_options.index(auto_display) if auto_display in prefix_options else 0
        prefix_chosen = st.selectbox(
            "字首", prefix_options, index=prefix_idx, key="fpo_prefix",
        )

    chosen_internal = internal_prefix(prefix_chosen)

    # ─── ★ v3.5 第 1 層防護:即時撈 Teable 編號(不靠 Refresh) ─
    live_po_numbers = get_live_po_numbers(table_url, headers)
    new_po_no = calc_next_po_number(live_po_numbers, chosen_internal)

    # 顯示編號 + Refresh 按鈕
    cols_no = st.columns([2, 0.6, 1, 1])
    with cols_no[0]:
        st.success(f"📋 新訂單編號:**`{new_po_no}`**")
    with cols_no[1]:
        if st.button("🔄", key="fpo_refresh_no", help="重新從 Teable 撈最新編號"):
            get_live_po_numbers(table_url, headers, force_refresh=True)
            st.rerun()
    with cols_no[2]:
        order_date_input = st.date_input("採購日期", value=date.today(), key="fpo_order_date")
    with cols_no[3]:
        is_revised = st.checkbox("REVISED", value=False, key="fpo_revised")

    # 顯示 cache 狀態(讓使用者放心)
    cache_t = st.session_state.get("_fpo_live_po_numbers_t", 0)
    if cache_t:
        age_sec = int(datetime.now().timestamp() - cache_t)
        st.caption(
            f"🟢 Teable 即時資料(共 {len(live_po_numbers)} 筆訂單,{age_sec} 秒前撈) · "
            "30 秒 cache 內不重撈 · 按 🔄 強制刷新"
        )

    cols_pic = st.columns([1, 3])
    with cols_pic[0]:
        purchase_responsible = st.text_input("負責採購", value="Amy", key="fpo_pic")

    # ─── Step 4 ─
    st.markdown("### Step 4. 每品項:工廠單價 + 工廠交期")

    # ★ v3.9.13: sentinel 只看 PDF identity(不再包含工廠/字首),
    # 換 PDF 才清 widget;切工廠/字首時保留 Sandy 已填的單價。
    step4_sentinel = f"{parsed.customer_po_no or ''}|{len(parsed.items)}"
    if st.session_state.get("_fpo_step4_sentinel") != step4_sentinel:
        # 新 PDF → 清掉 widget value 跟所有 prev_price cache
        for k in list(st.session_state.keys()):
            if k.startswith("fpo_fprice_") or k.startswith("_fpo_prev_price_"):
                try:
                    del st.session_state[k]
                except KeyError:
                    pass
        st.session_state["_fpo_step4_sentinel"] = step4_sentinel

    factory_unit_prices = {}
    factory_due_dates = {}
    current_po_prefix_for_step4 = display_prefix(chosen_internal)  # GC / ET / EW

    # ★ v3.9.13: debug 用 — 收集每個品項的查詢結果
    debug_prev_price_log = []

    for idx, it in enumerate(parsed.items):
        # ★ v3.9.13: prev_price cache key 包含 (idx, factory, prefix)
        # 切工廠/字首時自動重新計算,顯示更新的提示,但不會把 widget value 歸零
        prev_price_cache_key = (
            f"_fpo_prev_price_{idx}_{factory_short}_{current_po_prefix_for_step4}"
        )
        if prev_price_cache_key not in st.session_state:
            calc_result = fetch_previous_factory_price(
                orders, it.part_number, factory_short, current_po_prefix_for_step4
            )
            st.session_state[prev_price_cache_key] = calc_result
        prev_price = st.session_state[prev_price_cache_key]

        # debug log
        debug_prev_price_log.append({
            "idx": idx,
            "P/N": it.part_number,
            "Step3 工廠": factory_short,
            "Step3 字首": current_po_prefix_for_step4,
            "前批單價": prev_price if prev_price else "(沒找到)",
        })

        p_cols = st.columns([2, 1, 1, 2])
        with p_cols[0]:
            st.text_input("P/N", value=it.part_number, disabled=True, key=f"fpo_disp_pn_{idx}")
        with p_cols[1]:
            st.text_input("數量", value=str(it.quantity), disabled=True, key=f"fpo_disp_qty_{idx}")
        with p_cols[2]:
            # ★ v3.9.13: 只在 widget 還沒被建立過時用 prev_price 當 default
            #   (streamlit number_input 在 session_state 已存在時 value 參數會被忽略,
            #    所以 Sandy 編輯過的單價會被保留)
            widget_key = f"fpo_fprice_{idx}"
            if widget_key not in st.session_state:
                default_price = float(prev_price) if prev_price else 0.0
                # 顯式 set 進 session_state,讓 number_input 第一次看到就有值
                st.session_state[widget_key] = default_price
            f_price = st.number_input(
                "工廠單價",
                min_value=0.0,
                step=0.01,
                format="%.4f",
                key=widget_key,
                help=(
                    f"💡 前批單價: {prev_price}(切工廠/字首會重新計算這個值)"
                    if prev_price
                    else "(沒前批紀錄,需手動填)"
                ),
            )
            if prev_price:
                # 如果 widget value 跟 prev_price 不同,顯示一個提示
                if abs(float(f_price) - float(prev_price)) < 0.001:
                    st.caption(f"💡 前批單價: **{prev_price}**(已自動帶入)")
                else:
                    st.caption(
                        f"💡 前批單價: **{prev_price}** "
                        f"(目前單價 {f_price} 跟前批不同,Sandy 已手動編輯)"
                    )
        with p_cols[3]:
            default_due = date.today()
            if it.delivery_date:
                try:
                    default_due = datetime.strptime(it.delivery_date, "%Y-%m-%d").date()
                except ValueError:
                    pass
            f_due = st.date_input("工廠交期", value=default_due, key=f"fpo_fdue_{idx}")
        factory_unit_prices[it.part_number] = f_price
        factory_due_dates[it.part_number] = f_due

    # ★ v3.9.13: debug expander — 看 fetch_previous_factory_price 對每個品項的結果
    with st.expander("🔍 Debug: 前批單價查詢結果 (展開看 fetch_previous_factory_price 的詳情)", expanded=False):
        debug_df = pd.DataFrame(debug_prev_price_log)
        st.dataframe(debug_df, use_container_width=True, hide_index=True)
        # 順便顯示 orders DataFrame 的關鍵欄位
        if orders is not None and not orders.empty:
            st.caption(f"Teable orders 共 {len(orders)} 筆,欄位:")
            cols_debug = list(orders.columns)
            st.code(", ".join(cols_debug), language=None)
            # 對每個品項抽出該 P/N 在 orders 的紀錄
            for it in parsed.items:
                pn_target = str(it.part_number).strip()
                if "P/N" in orders.columns:
                    matches = orders[orders["P/N"].astype(str).str.strip() == pn_target]
                    st.write(f"**{pn_target}** 在 orders 找到 **{len(matches)}** 筆")
                    if len(matches) > 0:
                        # 顯示關鍵欄位
                        show_cols = [c for c in [
                            COL_GLOCOM_PO, "P/N", "工廠",
                            "Order Q'TY\n (PCS)", "Order Q'TY (PCS)", "Q'TY (PCS)",
                            "銷貨金額", "工廠下單日期",
                        ] if c in matches.columns]
                        if show_cols:
                            st.dataframe(
                                matches[show_cols].head(5),
                                use_container_width=True,
                                hide_index=True,
                            )
        else:
            st.warning("orders DataFrame 是空的(側邊欄可能要 Refresh)")

    factory_total = sum(
        factory_unit_prices.get(it.part_number, 0.0) * it.quantity
        for it in parsed.items
    )
    st.metric("工廠總金額", f"{factory_data.get('default_currency', 'NT$')} {factory_total:,.2f}")

    # ─── ★ v3.8.1: NRE 工程費獨立欄位 + 依字首/貨幣 ─
    # 規則:
    #   - 工廠貨幣依 factory_data.default_currency(NTD 台灣 / USD 大陸)
    #   - GC 字首 → NRE 平攤合併進主品項單價
    #   - ET/EW   → NRE 獨立一行(實質是給 99-9999 那行的工廠單價)
    #   - NRE 金額填 0 或 NA → 完全跳過
    
    # 工廠貨幣標準化:NT$ / NT → NTD,USD / US$ → USD
    raw_currency = factory_data.get("default_currency", "NTD")
    factory_currency = raw_currency.replace("$", "").strip().upper() or "NTD"
    if factory_currency == "NT":
        factory_currency = "NTD"
    elif factory_currency == "US":
        factory_currency = "USD"
    nre_mode = "merge" if chosen_internal == "G" else "separate"
    nre_mode_label = (
        f"GC 模式:平攤合併"
        if nre_mode == "merge"
        else f"{prefix_chosen} 模式:獨立一行"
    )
    
    st.markdown(f"### Step 4-1. NRE 工程費 — `{nre_mode_label}` · 工廠貨幣 `{factory_currency}`")

    # ★ v3.9.11: 換 PDF 時強制重置 NRE 為 "0"
    # 用 (customer_po_no + po_date + 品項數) 當 sentinel,變了就清 NRE session_state
    pdf_sentinel = f"{parsed.customer_po_no or ''}|{parsed.po_date or ''}|{len(parsed.items)}"
    if st.session_state.get("_fpo_pdf_sentinel", None) != pdf_sentinel:
        # 新 PDF → 清 NRE 欄位(streamlit text_input 的 widget value 殘留)
        st.session_state["fpo_nre_amount_str"] = "0"
        st.session_state["_fpo_pdf_sentinel"] = pdf_sentinel
    
    if nre_mode == "merge":
        st.caption(
            f"**GC 字首** → NRE 平攤到主品項工廠單價,規格欄底下加註,工廠 PDF 不印 NRE 那一行。  \n"
            f"**填 0 / NA / 留空 → 不處理 NRE**(跳過)"
        )
    else:
        st.caption(
            f"**{prefix_chosen} 字首** → NRE 保留獨立一行,工廠 PDF 印 NRE 品項。  \n"
            f"**填 0 / NA / 留空 → 不處理 NRE,客戶 PDF 上的 NRE 品項不會印**"
        )
    
    # 自動偵測 NRE 品項
    # 規則(從嚴):
    #   1. 料號 99-開頭 → NRE(VORNE 等客戶都用 99-XXXX)
    #   2. 料號完全等於 "NRE" → NRE
    #   3. 描述「整段是 NRE 相關」+ 數量為 1 → NRE
    #      只認:^NRE / ^1 Time NRE / ^Setup Fee / ^Tooling Fee / ^Engineering Fee 開頭
    #   ★ 不再用「描述含 Net List」做判斷 — Net List Test 常是 PCB 描述的一部分
    NRE_DESCRIPTION_PATTERNS = (
        "NRE", "1 TIME NRE", "1-TIME NRE",
        "SETUP FEE", "TOOLING FEE", "ENGINEERING FEE",
        "ONE-TIME FEE", "ONE TIME FEE",
    )
    
    def _is_nre_item(it):
        # 料號規則
        pn = (it.part_number or "").strip()
        if pn.startswith("99-"):
            return True
        if pn.upper() == "NRE":
            return True
        # 描述規則(只在數量為 1 才考慮 — 多顆通常不是 NRE)
        desc = (it.description or "").strip().upper()
        if it.quantity == 1:
            for pattern in NRE_DESCRIPTION_PATTERNS:
                if desc.startswith(pattern):
                    return True
            # 完全等於某些 keywords
            if desc in ("NRE", "1 TIME NRE", "SETUP", "TOOLING"):
                return True
        return False
    
    detected_nre_items = []
    main_items = []
    for it in parsed.items:
        if _is_nre_item(it):
            detected_nre_items.append(it)
        else:
            main_items.append(it)
    
    if detected_nre_items:
        st.info(
            f"💡 自動偵測到 {len(detected_nre_items)} 個 NRE 品項:"
            f" {', '.join(it.part_number for it in detected_nre_items)}"
            f"(料號 99-開頭 或 描述以 NRE/Setup Fee/Tooling Fee 開頭 + 數量=1)"
        )
    
    nre_amount = 0.0
    nre_target_pn = None
    has_nre = False
    
    # ─── 獨立 NRE 輸入欄位 ─
    nre_cols = st.columns([2, 2, 3])
    
    with nre_cols[0]:
        # ★ v3.9.9: NRE 金額預設「0」(不再依客戶 PDF 自動帶估算值,避免 Sandy 沒注意就送出)
        # GC / ET / EW 三種模式統一預設「0」,Sandy 真有 NRE 才手動填數字
        default_nre_str = "0"

        nre_input_str = st.text_input(
            f"NRE 金額 ({factory_currency})",
            value=default_nre_str,
            key="fpo_nre_amount_str",
            help=f"工廠開的 NRE 金額。預設 0(無 NRE)。要處理 NRE 才把 0 改成實際金額。",
            placeholder="例: 14000 / 0 / NA",
        )
        
        # 解析輸入:0 / NA / 空 → 不處理
        nre_input_clean = (nre_input_str or "").strip().upper()
        if nre_input_clean in ("", "0", "NA", "N/A", "NONE", "-"):
            has_nre = False
            nre_amount = 0.0
        else:
            try:
                nre_amount = float(nre_input_clean.replace(",", ""))
                if nre_amount > 0:
                    has_nre = True
            except ValueError:
                st.error(f"❌ NRE 金額格式錯誤:`{nre_input_str}` — 請填數字、0、或 NA")
                nre_amount = 0.0
                has_nre = False
    
    if has_nre and nre_mode == "merge":
        # GC 模式:選主品項
        with nre_cols[1]:
            if main_items:
                pn_options = [it.part_number for it in main_items]
                nre_target_pn = st.selectbox(
                    "NRE 對應到主品項 *",
                    options=pn_options,
                    key="fpo_nre_target",
                    help="NRE 金額會除以這個品項的數量,加到工廠單價",
                )
            else:
                st.warning("⚠️ 沒有主品項可以攤提 NRE")
        
        with nre_cols[2]:
            if nre_target_pn and nre_amount > 0:
                target_item = next(
                    (it for it in main_items if it.part_number == nre_target_pn),
                    None,
                )
                if target_item and target_item.quantity > 0:
                    nre_per_pcs = nre_amount / target_item.quantity
                    original_price = factory_unit_prices.get(nre_target_pn, 0.0)
                    new_price = original_price + nre_per_pcs
                    st.success(
                        f"📊 NRE {nre_amount:,.2f} {factory_currency} ÷ {target_item.quantity} pcs"
                        f" = {nre_per_pcs:.2f} {factory_currency}/pcs\n\n"
                        f"主品項 {nre_target_pn} 工廠單價:\n"
                        f"{factory_currency} {original_price:.2f} + {factory_currency} {nre_per_pcs:.2f}"
                        f" = **{factory_currency} {new_price:.2f}**"
                    )
        
        # 套到 factory_unit_prices
        if nre_target_pn and nre_amount > 0:
            target_item = next(
                (it for it in main_items if it.part_number == nre_target_pn),
                None,
            )
            if target_item and target_item.quantity > 0:
                nre_per_pcs = nre_amount / target_item.quantity
                # ★ v3.9.5: 先記住「原本的工廠單價」(沒攤 NRE),寫進 session_state
                # 這樣 build_po_context 才有辦法在 NRE 註記裡顯示「單價NTD102/pcs」
                original_unit_price = factory_unit_prices.get(nre_target_pn, 0.0)
                st.session_state["_fpo_nre_original_unit_price"] = original_unit_price
                
                factory_unit_prices[nre_target_pn] = (
                    factory_unit_prices.get(nre_target_pn, 0.0) + nre_per_pcs
                )
    
    elif has_nre and nre_mode == "separate":
        # ET/EW 模式:NRE 獨立顯示,自動把 NRE 金額填到 99-9999 那行的 factory_unit_prices
        with nre_cols[1]:
            if detected_nre_items:
                # 如果有多筆 NRE,選一筆對應(通常只 1 筆)
                if len(detected_nre_items) > 1:
                    nre_target_pn = st.selectbox(
                        "NRE 金額對應到品項 *",
                        options=[it.part_number for it in detected_nre_items],
                        key="fpo_nre_target_separate",
                    )
                else:
                    nre_target_pn = detected_nre_items[0].part_number
                    st.text_input(
                        "NRE 品項",
                        value=nre_target_pn,
                        disabled=True,
                        key="fpo_nre_target_separate_display",
                    )
            else:
                st.warning(
                    "⚠️ 客戶 PDF 沒偵測到 NRE 品項。"
                    "如要保留 NRE 一行,請在 Step 4 主品項清單裡手動處理。"
                )
        
        with nre_cols[2]:
            if nre_target_pn and nre_amount > 0:
                st.success(
                    f"📊 NRE 品項 {nre_target_pn}\n\n"
                    f"工廠單價: **{factory_currency} {nre_amount:,.2f}**\n\n"
                    f"工廠 PDF 上會單獨印一行,單價 = 金額 = {nre_amount:,.2f} {factory_currency}"
                )
        
        # 把 NRE 金額直接填到 99-9999 那行
        if nre_target_pn and nre_amount > 0:
            factory_unit_prices[nre_target_pn] = nre_amount

    # 儲存 NRE 設定
    st.session_state["_fpo_nre_settings"] = {
        "mode": nre_mode,
        "has_nre": has_nre,
        "nre_amount": nre_amount,
        "nre_target_pn": nre_target_pn,
        "detected_nre_pns": [it.part_number for it in detected_nre_items],
        "factory_currency": factory_currency,
        # ★ v3.9.5: 主品項原本工廠單價(攤 NRE 之前),用來顯示在 PDF 的「單價NTD102/pcs」
        "original_unit_price": st.session_state.get("_fpo_nre_original_unit_price", 0.0),
    }

    # ─── Step 5: 每品項規格(3 Tab) ─
    st.markdown("### Step 5. 每品項產品規格")
    st.caption("每個品項各有 3 種輸入方式可選。規格會以「一列」印在工廠 PO 的『產品規格』欄。")
    
    # ★ v3.8.1: NRE 品項是否跳過 Step 5 規格輸入
    # 跟 build_po_context 用同樣的 skip 邏輯,確保不會印的品項才跳過
    nre_settings_for_step5 = st.session_state.get("_fpo_nre_settings", {}) or {}
    _mode = nre_settings_for_step5.get("mode", "merge")
    _has_nre = nre_settings_for_step5.get("has_nre", False)
    skip_nre_in_step5 = (
        (_mode == "merge") or                       # GC 任何情況都不在 Step 5 顯示 NRE
        (_mode == "separate" and not _has_nre)     # ET/EW 沒填 NRE 也跳過
    )
    skip_nre_pns_step5 = (
        set(nre_settings_for_step5.get("detected_nre_pns", []))
        if skip_nre_in_step5
        else set()
    )

    item_specs = {}
    # ★ v3.9.11: 把 Step 3 選的字首傳到每個品項的規格輸入,讓 default 按字首抓
    current_po_prefix_for_items = display_prefix(chosen_internal)  # "GC" / "ET" / "EW"
    for idx, it in enumerate(parsed.items):
        if it.part_number in skip_nre_pns_step5:
            with st.container(border=True):
                if _mode == "merge" and _has_nre:
                    msg = f"💡 品項 `{it.part_number}` 是 NRE,已合併到主品項,不會印工廠 PDF,跳過規格填寫"
                elif _mode == "merge":
                    msg = f"💡 品項 `{it.part_number}` 是 NRE,GC 模式且未填 NRE 金額,不會印工廠 PDF,跳過規格填寫"
                else:
                    msg = f"💡 品項 `{it.part_number}` 是 NRE,未填 NRE 金額,不會印工廠 PDF,跳過規格填寫"
                st.caption(msg)
            item_specs[it.part_number] = ""
            continue
        spec_text = render_spec_input_for_item(
            idx, it, orders,
            current_po_prefix=current_po_prefix_for_items,  # ★ v3.9.11
        )
        item_specs[it.part_number] = spec_text

    # ─── Step 6: 產出按鈕區(產 PDF / 寫回 Teable / 產 PI / 清除)─
    st.divider()
    st.markdown("### Step 6. 產出 / 寫回")
    st.caption("**產 PDF** / **寫回 Teable** / **產 PI** 是分開的。建議先產 PDF 預覽 → OK 再寫回 Teable / 產 PI。")

    # ★ v3.9.15: 模板診斷工具 — 為了 debug PDF 第二品項被吞 bug
    with st.expander(
        "🔧 PO_GLOCOM.docx 模板診斷工具(展開後可下載模板給 Claude 看)",
        expanded=False,
    ):
        st.warning(
            "**❗ PDF 只印一筆品項 bug 必看** — 程式已經送 2 筆品項給模板"
            "(看上面「產 PDF 前確認」expander 證實),但 docx 模板 jinja2 迴圈寫錯,"
            "只展開了第一筆。這個 bug 必須修模板,**請按下面按鈕下載模板給 Claude**。"
        )

        # 找 docx 模板路徑
        template_paths_to_try = [
            Path("templates/PO_GLOCOM.docx"),
            Path("/mount/src/glocom-wip-v36/templates/PO_GLOCOM.docx"),  # Streamlit Cloud
            Path(__file__).parent / "templates" / "PO_GLOCOM.docx",
        ]
        template_path = None
        for p in template_paths_to_try:
            if p.exists():
                template_path = p
                break

        if not template_path:
            st.error(
                "⚠️ 找不到 templates/PO_GLOCOM.docx。"
                "請從 GitHub 下載 → "
                "https://github.com/jonah58w/glocom-wip-v36/raw/main/templates/PO_GLOCOM.docx"
            )
        else:
            st.caption(f"找到模板:`{template_path}`")
            cols_diag = st.columns(2)
            with cols_diag[0]:
                # 一鍵下載當前 Streamlit Cloud 上跑的 docx 模板
                try:
                    with open(template_path, "rb") as f:
                        st.download_button(
                            "📥 下載 PO_GLOCOM.docx 給 Claude 看",
                            f.read(),
                            file_name="PO_GLOCOM.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            type="primary",
                            key="dl_template_for_claude",
                        )
                except Exception as e:
                    st.error(f"讀模板失敗:{e}")

            with cols_diag[1]:
                show_structure = st.button(
                    "🔍 顯示模板結構(table/row/cell)",
                    use_container_width=True,
                    key="show_template_structure",
                )

            if show_structure:
                try:
                    from docx import Document
                except ImportError:
                    st.error("python-docx 未安裝")
                else:
                    try:
                        doc = Document(str(template_path))
                        st.write(f"**Tables 數**: {len(doc.tables)}")
                        diag_lines = [f"# Template: {template_path}"]
                        diag_lines.append(f"Tables: {len(doc.tables)}\n")

                        for t_idx, table in enumerate(doc.tables):
                            diag_lines.append(
                                f"\n## Table {t_idx} — "
                                f"{len(table.rows)} rows × {len(table.columns)} cols"
                            )
                            for r_idx, row in enumerate(table.rows):
                                row_height = ""
                                try:
                                    if row.height:
                                        row_height = f" (height: {row.height})"
                                except Exception:
                                    pass
                                diag_lines.append(f"\n  ### Row {r_idx}{row_height}")
                                for c_idx, cell in enumerate(row.cells):
                                    cell_text = cell.text
                                    diag_lines.append(
                                        f"    Cell {c_idx}: {repr(cell_text[:300])}"
                                    )
                                    # 偵測 jinja2 tag
                                    if "{%" in cell_text or "{{" in cell_text:
                                        tags = []
                                        if "{%tr for" in cell_text or "{% tr for" in cell_text:
                                            tags.append("✅ 有 tr-for")
                                        if "{%tr endfor" in cell_text or "{% tr endfor" in cell_text:
                                            tags.append("✅ 有 tr-endfor")
                                        if "{% for item" in cell_text and "{%tr" not in cell_text and "{% tr" not in cell_text:
                                            tags.append("⚠️ 用了 {% for %} 但沒 tr 前綴!")
                                        if "{{ item." in cell_text or "{{item." in cell_text:
                                            tags.append("📦 含 item placeholder")
                                        if tags:
                                            diag_lines.append(f"      → {' | '.join(tags)}")

                        text_output = "\n".join(diag_lines)
                        st.code(text_output, language="text")
                        st.download_button(
                            "💾 下載診斷報告 (.txt)",
                            text_output,
                            file_name="po_glocom_template_diagnostic.txt",
                            key="dl_diag_report",
                        )
                    except Exception as e:
                        st.error(f"診斷失敗:{e}")
                        with st.expander("詳細錯誤"):
                            st.code(traceback.format_exc())

    # ★ v3.9.14: 顯示「PDF 即將使用的工廠/字首/編號」— Sandy 一眼確認
    with st.container(border=True):
        st.markdown("##### 📋 PDF 即將使用的設定(按下「產 PDF」會用這些值)")
        preview_cols = st.columns(4)
        preview_cols[0].markdown(f"**工廠**  \n`{factory_short}`")
        preview_cols[1].markdown(f"**字首**  \n`{prefix_chosen}` / {issuing_company}")
        preview_cols[2].markdown(f"**編號**  \n`{new_po_no}`")
        preview_cols[3].markdown(f"**合計**  \n`{factory_data.get('default_currency', 'NT$')} {factory_total:,.2f}`")
        # 顯示工廠基本資料(模板會用)
        st.caption(
            f"工廠資料:**{factory_data.get('factory_name', '(未設定)')}** | "
            f"{factory_data.get('address', '(未設定地址)')} | "
            f"{factory_data.get('phone', '(未設定電話)')}"
        )

    # ★ v3.9.14: 如果有上次產的 PDF 快照,且跟當前 Step 3 不同 → 顯示警告
    last_snap = st.session_state.get("_fpo_last_pdf_snapshot")
    if last_snap:
        is_stale = (
            last_snap.get("factory_short") != factory_short
            or last_snap.get("prefix") != prefix_chosen
            or last_snap.get("po_no") != new_po_no
        )
        if is_stale:
            st.warning(
                f"⚠️ **上次產的 PDF (`{last_snap['po_no']}`) 跟現在的設定不同** —  \n"
                f"上次:`{last_snap['factory_short']}` / `{last_snap['prefix']}` / `{last_snap['po_no']}`  \n"
                f"現在:`{factory_short}` / `{prefix_chosen}` / `{new_po_no}`  \n"
                f"**要套新設定請按「📄 產 PDF」重新產一次**。"
            )
    
    # ★ v3.9.7: PI 設定區(只在客戶是 USD 客戶時顯示)
    # 簡化:Sandy 自己決定要不要產 PI,Bank fee 可改
    pi_settings_expander = st.expander("📑 PI 設定(產 PI 前可調整)", expanded=False)
    with pi_settings_expander:
        cols_pi = st.columns(3)
        with cols_pi[0]:
            pi_bank_fee = st.number_input(
                "Bank fee (USD)",
                min_value=0.0, value=45.0, step=5.0,
                key="fpo_pi_bank_fee",
                help="預設 $45,可改",
            )
        with cols_pi[1]:
            pi_terms = st.text_input(
                "Terms",
                value="Exwork Taiwan (US$ Exwork Taiwan)",
                key="fpo_pi_terms",
            )
        with cols_pi[2]:
            pi_to_country = st.text_input(
                "To Country",
                value="USA",
                key="fpo_pi_to_country",
                help="客戶所在國,例如 USA / Japan",
            )
    
    btn_cols = st.columns([2, 2, 2, 1])
    with btn_cols[0]:
        do_pdf = st.button(
            "📄 產 PDF", type="primary",
            use_container_width=True, key="fpo_do_pdf",
        )
    with btn_cols[1]:
        do_writeback = st.button(
            "💾 寫回 Teable", type="secondary",
            use_container_width=True, key="fpo_do_writeback",
        )
    with btn_cols[2]:
        do_pi = st.button(
            "📑 產 PI", type="secondary",
            use_container_width=True, key="fpo_do_pi",
            help="產 Proforma Invoice(給客戶看的英文 PI,USD 報價)",
        )
    with btn_cols[3]:
        if st.button("🗑️ 全部清除", use_container_width=True, key="fpo_clear_bottom"):
            _reset_form()
            st.rerun()

    if not (do_pdf or do_writeback or do_pi):
        st.stop()

    # ─── 共用驗證 ─
    # ★ v3.8.1: NRE 跳過邏輯
    nre_settings_v = st.session_state.get("_fpo_nre_settings", {}) or {}
    _v_mode = nre_settings_v.get("mode", "merge")
    _v_has_nre = nre_settings_v.get("has_nre", False)
    _v_skip = (
        (_v_mode == "merge") or
        (_v_mode == "separate" and not _v_has_nre)
    )
    skip_nre_pns_v = (
        set(nre_settings_v.get("detected_nre_pns", []))
        if _v_skip
        else set()
    )
    
    zero_price_pns = [
        pn for pn, price in factory_unit_prices.items()
        if price == 0 and pn not in skip_nre_pns_v
    ]
    if zero_price_pns:
        st.warning(f"⚠️ 工廠單價為 0 的品項:{', '.join(zero_price_pns)}")

    empty_spec_pns = [
        pn for pn, sp in item_specs.items()
        if not (sp or "").strip() and pn not in skip_nre_pns_v
    ]
    if empty_spec_pns:
        st.error(f"❌ 以下品項規格沒填:**{', '.join(empty_spec_pns)}**。"
                 "規格欄會印空白,請回 Step 5 填好再產 PDF / 寫回 Teable。")
        st.stop()

    # ─── ★ v3.5 第 2 層防護:產 PDF 前撞號檢查 ─
    if do_pdf:
        with st.spinner("最後檢查 Teable 編號..."):
            fresh_po_set = get_live_po_numbers(table_url, headers, force_refresh=True)
        
        if new_po_no in fresh_po_set:
            # 撞了!算新的並擋下
            corrected_po = calc_next_po_number(fresh_po_set, chosen_internal)
            st.error(
                f"❌ **編號撞號警告!** "
                f"剛剛偵測到 Teable 上已經有 **`{new_po_no}`**(可能是別人在你建單期間手動 key 進去)。\n\n"
                f"➡️ 系統建議改用 **`{corrected_po}`**。請按 🔄 重新整理頁面後再試一次。"
            )
            st.stop()

    if do_writeback:
        with st.spinner("最後檢查 Teable 編號..."):
            fresh_po_set = get_live_po_numbers(table_url, headers, force_refresh=True)
        
        if new_po_no in fresh_po_set:
            # ─── ★ v3.5 第 3 層:寫入鎖,撞了就拒絕 ─
            corrected_po = calc_next_po_number(fresh_po_set, chosen_internal)
            st.error(
                f"❌ **寫入鎖觸發!** "
                f"Teable 上已經有 **`{new_po_no}`**,拒絕寫入避免覆蓋。\n\n"
                f"➡️ 請按 🔄 重新整理編號(會跳到 **`{corrected_po}`**)後重試。"
            )
            st.stop()

    # ─── 產 PDF ─
    if do_pdf:
        po_ctx = build_po_context_from_new_order(
            new_po_no=new_po_no, parsed=parsed, factory_data=factory_data,
            factory_short=factory_short, issuing_company=issuing_company,
            factory_unit_prices=factory_unit_prices, factory_due_dates=factory_due_dates,
            item_specs=item_specs, purchase_responsible=purchase_responsible,
            order_date=order_date_input, is_revised=is_revised,
            nre_settings=st.session_state.get("_fpo_nre_settings"),  # ★ v3.8
        )

        # ★ v3.9.13: 產 PDF 前先顯示 items 摘要 — 確認 docx 模板會收到的所有品項
        with st.expander(
            f"🔍 產 PDF 前確認 — items list 共 **{len(po_ctx.get('items', []))}** 個品項",
            expanded=True,
        ):
            items_summary = []
            for it_ctx in po_ctx.get("items", []):
                spec_preview = (it_ctx.get("spec_text") or "").replace("\u200B", "").split("\n")[0][:50]
                items_summary.append({
                    "P/N": it_ctx["part_number"],
                    "數量": it_ctx["quantity"],
                    "工廠單價": it_ctx["unit_price"],
                    "小計": it_ctx["amount"],
                    "規格首行": spec_preview,
                })
            if items_summary:
                st.dataframe(
                    pd.DataFrame(items_summary),
                    use_container_width=True,
                    hide_index=True,
                )
                expected_total = sum(float(it.get("amount", 0)) for it in po_ctx["items"])
                st.caption(
                    f"⚠️ 如果 PDF 上的品項列數跟這裡不一致(譬如這裡有 {len(items_summary)} 筆 "
                    f"但 PDF 只印 1 筆),問題在 `templates/PO_GLOCOM.docx` 的 "
                    f"`{{% tr for item in items %}}` 迴圈,需要修模板。  \n"
                    f"預期合計:`{po_ctx.get('currency', 'NT$')} {expected_total:,.2f}`"
                )
            else:
                st.error("⚠️ items list 是空的!產 PDF 會失敗。")

        try:
            from core.pdf_generator import generate_po_files
            with st.spinner("產生 docx + PDF..."):
                result = generate_po_files(po_ctx)
            docx_path = result.get("docx_path")
            pdf_path = result.get("pdf_path")
            if result.get("error"):
                st.warning(result["error"])
        except FileNotFoundError as e:
            if "PO_EUSWAY" in str(e):
                st.error(
                    "❌ EUSWAY 模板未上傳:`templates/PO_EUSWAY.docx` 不存在。"
                )
            else:
                st.error(f"❌ 產 PDF 失敗:{e}")
            st.stop()
        except Exception as e:
            st.error(f"❌ 產 PDF 失敗:{e}")
            with st.expander("錯誤詳情"):
                st.code(traceback.format_exc())
            st.stop()

        st.success(f"✅ PDF / DOCX 已產生(訂單編號 {new_po_no})")

        # ★ v3.9.14: 記錄這份 PDF 是用什麼工廠/字首產的(避免 Sandy 切工廠後看 PDF 誤以為是新版)
        st.session_state["_fpo_last_pdf_snapshot"] = {
            "po_no": new_po_no,
            "factory_short": factory_short,
            "factory_name": factory_data.get("factory_name", ""),
            "prefix": prefix_chosen,
            "issuing_company": issuing_company,
            "items_count": len(po_ctx.get("items", [])),
            "total_amount": po_ctx.get("total_amount", 0),
            "currency": po_ctx.get("currency", "NT$"),
        }

        # 顯示 PDF 快照資訊(讓 Sandy 一眼看到 PDF 是用什麼產的)
        snap = st.session_state["_fpo_last_pdf_snapshot"]
        st.info(
            f"📋 **這份 PDF 的快照** —  \n"
            f"• 編號:**`{snap['po_no']}`**  \n"
            f"• 工廠:**{snap['factory_short']}**({snap['factory_name']})  \n"
            f"• 字首/發出公司:**{snap['prefix']} / {snap['issuing_company']}**  \n"
            f"• 品項數:**{snap['items_count']}** 筆  \n"
            f"• 合計:**{snap['currency']} {snap['total_amount']:,.2f}**\n\n"
            f"⚠️ 之後如果切了 Step 3 工廠/字首,要再按一次「📄 產 PDF」才會用新值產。"
        )

        d_cols = st.columns(2)
        if docx_path and Path(docx_path).exists():
            with d_cols[0]:
                with open(docx_path, "rb") as f:
                    st.download_button(
                        "📥 下載 DOCX", data=f.read(),
                        file_name=Path(docx_path).name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True, key="fpo_dl_docx",
                    )
        if pdf_path and Path(pdf_path).exists():
            with d_cols[1]:
                with open(pdf_path, "rb") as f:
                    st.download_button(
                        "📥 下載 PDF", data=f.read(),
                        file_name=Path(pdf_path).name, mime="application/pdf",
                        use_container_width=True, key="fpo_dl_pdf",
                    )

        st.info("📝 PDF 確認 OK 後,可按上方「💾 寫回 Teable」把資料寫進主表。")

    # ─── 寫回 Teable ─
    if do_writeback:
        # ★ v3.8.1: NRE 設定
        nre_settings = st.session_state.get("_fpo_nre_settings", {}) or {}
        nre_mode_wb = nre_settings.get("mode", "merge")
        has_nre_wb = nre_settings.get("has_nre", False)
        nre_amount_wb = float(nre_settings.get("nre_amount", 0) or 0)
        nre_target_pn_wb = nre_settings.get("nre_target_pn")
        detected_nre_pns_wb = set(nre_settings.get("detected_nre_pns", []))
        factory_currency_wb = nre_settings.get("factory_currency", "NTD")
        
        apply_merge_wb = (nre_mode_wb == "merge" and has_nre_wb)
        skip_nre_in_writeback = (
            (nre_mode_wb == "merge") or
            (nre_mode_wb == "separate" and not has_nre_wb)
        )
        
        records_fields = []
        for it in parsed.items:
            # 跳過 NRE 品項的條件:跟 build_po_context 一致
            if it.part_number in detected_nre_pns_wb and skip_nre_in_writeback:
                continue
            
            f_price = factory_unit_prices.get(it.part_number, 0.0)
            f_due = factory_due_dates.get(it.part_number)
            f_due_str = f_due.strftime("%Y/%m/%d") if f_due else ""
            order_date_str = order_date_input.strftime("%Y/%m/%d")
            cust_date_str = parsed.po_date or order_date_str
            spec_text = (item_specs.get(it.part_number, "") or "").strip()
            
            # ★ v3.9.5: GC merge + has_nre 時主品項加註(新格式含原單價)
            #   寫進 Teable「客戶要求注意事項」欄位用單行版本(逗號分隔,不換行)
            if apply_merge_wb and it.part_number == nre_target_pn_wb and nre_amount_wb > 0 and it.quantity > 0:
                nre_per_pcs = nre_amount_wb / it.quantity
                cur = factory_currency_wb if factory_currency_wb != "NTD" else "NTD"
                # 取原單價(攤 NRE 之前的)
                orig_price = float(nre_settings.get("original_unit_price", 0.0) or 0.0)
                nre_note = (
                    f"(內含 單價{cur}{orig_price:.0f}/pcs"
                    f" + NRE: {cur}{nre_amount_wb:,.0f}; {cur}{nre_per_pcs:.0f}/pcs)"
                )
                if spec_text:
                    spec_text = f"{spec_text}\n{nre_note}"
                else:
                    spec_text = nre_note
            
            # GC merge + has_nre 時主品項接單金額補上 NRE 那行的客戶金額
            customer_amount = round(it.amount, 2)
            if apply_merge_wb and it.part_number == nre_target_pn_wb:
                for nre_it in parsed.items:
                    if nre_it.part_number in detected_nre_pns_wb:
                        customer_amount = round(it.amount + nre_it.amount, 2)
                        break

            records_fields.append({
                COL_GLOCOM_PO: new_po_no,
                "客戶": parsed.customer_name,
                "PO#": parsed.customer_po_no,
                "P/N": it.part_number,
                "Order Q'TY\n (PCS)": it.quantity,
                "工廠": factory_short,
                "工廠交期": f_due_str,
                "Ship date": f_due_str,
                "銷貨金額": round(f_price * it.quantity, 2),
                "接單金額": customer_amount,
                "客戶要求注意事項": spec_text,
                "工廠下單日期": order_date_str,
                "客戶下單日期": cust_date_str,
            })

        with st.spinner(f"寫入 {len(records_fields)} 筆 Teable record..."):
            wb_result = create_teable_records(table_url, headers, records_fields)
        
        # ★ v3.9.6: 偵錯訊息 — 確保任何狀況都有提示
        if not records_fields:
            st.error(
                f"❌ **沒有任何 record 可寫入 Teable**(records_fields 是空的)\n\n"
                f"可能原因:NRE 偵測把所有品項都當 NRE 跳過。\n\n"
                f"偵錯資訊:\n"
                f"- parsed.items 共 {len(parsed.items)} 個品項: {[it.part_number for it in parsed.items]}\n"
                f"- detected_nre_pns: {sorted(detected_nre_pns_wb)}\n"
                f"- nre_mode: {nre_mode_wb}, has_nre: {has_nre_wb}, skip_nre_in_writeback: {skip_nre_in_writeback}\n\n"
                f"➡️ 請按「🗑 全部清除」重新建單,或回 Step 4-1 確認 NRE 偵測。"
            )
            st.stop()

        if wb_result["success"] > 0:
            st.success(f"✅ 已寫入 Teable {wb_result['success']} 筆 record(訂單編號 {new_po_no})")
            # 寫成功後清掉 cache,下次進來會重撈
            for k in ["_fpo_live_po_numbers", "_fpo_live_po_numbers_t"]:
                if k in st.session_state:
                    del st.session_state[k]
            st.info("⚠️ 請按側邊欄 **Refresh** 才會看到主表更新。")
            
            # ─── ★ v3.6: 順便寫進 spec_history.json (GitHub) ─
            try:
                from core.spec_history_writer import (
                    append_multiple_spec_history,
                    is_github_writer_available,
                )
                if is_github_writer_available():
                    sh_items = []
                    for it in parsed.items:
                        sh_items.append({
                            "part_number": it.part_number,
                            "spec_text": (item_specs.get(it.part_number, "") or "").strip(),
                            "po_no": new_po_no,
                            "factory": factory_short,
                            "date_str": order_date_input.strftime("%Y-%m-%d"),
                        })
                    
                    with st.spinner("更新規格歷史庫到 GitHub..."):
                        sh_success, sh_fail, sh_msgs = append_multiple_spec_history(sh_items)
                    
                    if sh_success > 0 and sh_fail == 0:
                        st.success(
                            f"✅ 規格歷史庫已即時更新到 GitHub ({sh_success} 筆) — "
                            "下一張同料號訂單會自動帶入此規格"
                        )
                    elif sh_fail > 0:
                        st.warning(
                            f"⚠️ 規格歷史庫 GitHub 更新部分失敗(成功 {sh_success}, 失敗 {sh_fail})"
                        )
                    
                    with st.expander("📜 規格歷史庫更新明細"):
                        for m in sh_msgs:
                            st.text(m)
                else:
                    st.caption(
                        "💡 提示:設定 GitHub Token 後可即時更新規格歷史庫。"
                        "暫時不影響使用 — 排程腳本(13:30)仍會處理 ERP 出單。"
                    )
            except Exception as e:
                st.warning(f"⚠️ 規格歷史庫更新失敗(不影響 Teable 主表): {e}")
        
        if wb_result["failed"] > 0:
            st.error(f"❌ 寫入失敗 {wb_result['failed']} 筆")
            with st.expander("失敗原因"):
                for err in wb_result["errors"]:
                    st.text(err)
        elif wb_result["success"] == 0:
            # ★ v3.9.6: 寫入既沒成功也沒失敗計數,可能是 API 回應異常
            st.error(
                f"❌ **寫入結果異常**(success=0, failed=0)\n\n"
                f"偵錯資訊:\n"
                f"- 嘗試寫入 {len(records_fields)} 筆\n"
                f"- API 回應沒有正常 records 列表\n\n"
                f"➡️ 請檢查 Teable API token 是否有效,或開瀏覽器 console 看詳細錯誤。"
            )

    # ─── ★ v3.9.7: 產 PI(Proforma Invoice)─
    if do_pi:
        try:
            from core.pi_generator import generate_pi_files
        except ImportError:
            try:
                from pi_generator import generate_pi_files
            except ImportError:
                st.error("❌ 找不到 pi_generator.py,請先部署。")
                st.stop()
        
        # 讀客戶檔(data/customers.json)抓客戶完整資訊
        import json as _json
        customers_data = {}
        try:
            customer_file_paths = [
                Path("data/customers.json"),
                Path("/mount/src/glocom-wip-v36/data/customers.json"),  # Streamlit Cloud
                Path(__file__).parent / "data" / "customers.json",
            ]
            for cust_path in customer_file_paths:
                if cust_path.exists():
                    with open(cust_path, "r", encoding="utf-8") as f:
                        customers_data = _json.load(f)
                    break
        except Exception as e:
            st.warning(f"⚠️ 讀取客戶檔失敗: {e}(會用空白資訊產 PI)")
        
        # 找對應客戶資料(用 customer 名稱當 key)
        customer_key = (parsed.customer_name or "").strip().upper()
        customer_info = customers_data.get(customer_key, {})
        
        if not customer_info:
            st.warning(
                f"⚠️ 客戶檔裡沒有 **{customer_key}** 的資料,PI 上客戶欄位會空白。\n\n"
                f"請在 GitHub `data/customers.json` 加入此客戶。"
            )
        
        # 算 PI 主號(去掉 -01)
        pi_invoice_no = new_po_no.split("-")[0] if "-" in new_po_no else new_po_no
        
        # 從 NRE 設定 + parsed.items 組 PI items
        nre_settings_pi = st.session_state.get("_fpo_nre_settings", {}) or {}
        detected_nre_pns_pi = set(nre_settings_pi.get("detected_nre_pns", []))
        
        pi_items = []
        for it in parsed.items:
            if it.part_number in detected_nre_pns_pi:
                # NRE 品項用 Setup & Tooling Charge 形式進 PI
                pi_items.append({
                    "item_no": nre_settings_pi.get("nre_target_pn") or it.part_number,
                    "description": "Setup & Tooling Charge",
                    "quantity": 1,
                    "quantity_unit": "set",
                    "unit_price": float(it.unit_price),
                    "amount": float(it.amount),
                    "is_setup": True,
                })
            else:
                # PCB 主品項 - 用 PCB 規格英文化或直接用解析的描述
                desc = (it.description or "").strip()
                if not desc:
                    desc = "Bare Printed Circuit Board"
                pi_items.append({
                    "item_no": it.part_number,
                    "description": desc,
                    "quantity": int(it.quantity),
                    "quantity_unit": "pcs",
                    "unit_price": float(it.unit_price),
                    "amount": float(it.amount),
                    "is_setup": False,
                })
        
        # 組 PI context
        from datetime import date as date_cls
        pi_ctx = {
            "invoice_no": pi_invoice_no,
            "po_no": new_po_no,
            "customer_po_no": parsed.customer_po_no or "",
            "customer_short": parsed.customer_name or "",
            "customer_code": customer_info.get("code", ""),
            "date": date_cls.today(),
            "contact_person": customer_info.get("contact_person", ""),
            "customer_name": customer_info.get("name_full", parsed.customer_name or ""),
            "customer_address1": customer_info.get("address_line1", ""),
            "customer_address2": customer_info.get("address_line2", ""),
            "customer_tel": customer_info.get("tel", ""),
            "customer_fax": customer_info.get("fax", ""),
            "shipment_text": customer_info.get(
                "default_shipment_text",
                "Since the T/T payment is received and WG approved, we will confirm the ship date then",
            ),
            "from_country": "Taiwan",
            "to_country": st.session_state.get("fpo_pi_to_country") or customer_info.get("country", "USA"),
            "terms_text": st.session_state.get("fpo_pi_terms") or customer_info.get(
                "default_terms", "Exwork Taiwan (US$ Exwork Taiwan)"
            ),
            "items": pi_items,
            "bank_fee": float(st.session_state.get("fpo_pi_bank_fee", 45.0)),
            "currency_symbol": "$",
        }
        
        st.divider()
        st.markdown("#### 📑 產 PI 結果")
        with st.spinner("產生 Proforma Invoice..."):
            try:
                pi_result = generate_pi_files(pi_ctx)
                
                if pi_result.get("error"):
                    st.warning(f"⚠️ {pi_result['error']}")
                
                # 提供下載
                docx_path = pi_result.get("docx_path")
                pdf_path = pi_result.get("pdf_path")
                
                cols_dl = st.columns(2)
                if pdf_path and Path(pdf_path).exists():
                    with cols_dl[0]:
                        with open(pdf_path, "rb") as f:
                            st.download_button(
                                "⬇ 下載 PI PDF",
                                f.read(),
                                file_name=Path(pdf_path).name,
                                mime="application/pdf",
                                use_container_width=True,
                                type="primary",
                                key="pi_dl_pdf",
                            )
                if docx_path and Path(docx_path).exists():
                    with cols_dl[1]:
                        with open(docx_path, "rb") as f:
                            st.download_button(
                                "⬇ 下載 PI DOCX(可編輯)",
                                f.read(),
                                file_name=Path(docx_path).name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True,
                                key="pi_dl_docx",
                            )
                
                st.success(
                    f"✅ Proforma Invoice **{pi_invoice_no}** 已產出\n\n"
                    f"- 客戶: {pi_ctx['customer_name']}\n"
                    f"- 客戶 PO: {pi_ctx['customer_po_no']}\n"
                    f"- 品項數: {len(pi_items)}\n"
                    f"- Bank fee: ${pi_ctx['bank_fee']}\n"
                )
                
                # 預覽 PI items
                with st.expander("📋 PI 品項清單(預覽)", expanded=False):
                    df_pi = pd.DataFrame([
                        {
                            "Item No.": it["item_no"],
                            "Description": it["description"][:60] + ("..." if len(it["description"]) > 60 else ""),
                            "Qty": f"{it['quantity']} {it['quantity_unit']}",
                            "Unit Price": f"${it['unit_price']}",
                            "Amount": f"${it['amount']}",
                        }
                        for it in pi_items
                    ])
                    st.dataframe(df_pi, use_container_width=True, hide_index=True)
                    if pi_ctx['bank_fee'] > 0:
                        st.caption(f"+ Bank fee: ${pi_ctx['bank_fee']}")
                    total = sum(float(it['amount']) for it in pi_items) + pi_ctx['bank_fee']
                    st.caption(f"**Total: ${total:,.2f}**")
            except Exception as e:
                st.error(f"❌ 產 PI 失敗: {e}")
                import traceback
                with st.expander("詳細錯誤"):
                    st.code(traceback.format_exc())
